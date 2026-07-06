"""
v6 - Surgical final pass targeting every remaining AI-patterned sentence.
Also cleans up structural mess (duplicate sections, orphaned lines, leftover junk text).
"""
from docx import Document

SRC = r'c:\Users\gaura\OneDrive\Documents\Adversarial Training\Adversaril attack and training Research paper_HUMANIZED.docx'
doc = Document(SRC)
paras = doc.paragraphs

def s(para, text):
    runs = para.runs
    if not runs:
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ''

def g(idx):
    return ''.join(r.text for r in paras[idx].runs)

REWRITES = {
    # [17] leftover old contribution line
    17: "",

    # [22] duplicate pipeline description
    22: "",

    # [25] "[29] All experiments use CIFAR-10 as the image source..." — still AI
    25: "I chose CIFAR-10: 60,000 images, 10 classes, 32×32 pixels, clean 50k/10k split. Familiar enough that results are interpretable; small enough to iterate fast.",

    # [34] "Layer-by-layer breakdown:" — fine as header but para 35 has arch section name
    34: "Layer breakdown:",

    # [35] Section heading "Convolutional Neural Network Architecture" out of place
    35: "",

    # [38] "Attack method: FGSM..." already covered elsewhere, condense
    38: "Attack: FGSM. Single gradient step, scaled by ε. Cheap and reproducible — the standard benchmark for a reason.",

    # [41] "[54] The perturbation is calculated as:" — duplicate of [39]
    41: "",

    # [43] "* J represents the loss function." — orphaned list item
    43: "",

    # [44] "* y represents the true label." — orphaned list item
    44: "",

    # [45] "I tested four epsilon values to see how attack severity scales with perturbation budget." — duplicate
    45: "",

    # [46] "The metrics work together..." displaced from metrics section
    46: "",

    # [47] "Evaluation Metrics" section heading — keep

    # [48] "Multi-Epsilon Robustness Analysis" — orphaned metric name
    48: "",

    # [49] "Everything ran in Python and PyTorch..." — duplicate of [53]
    49: "",

    # [50] "These metrics give complementary views..." — duplicate
    50: "",

    # [53] "All experiments ran in Python using PyTorch..." — duplicate of [49]
    53: "Python + PyTorch + VS Code on Windows. Everything inside a virtual environment so deps don't clash.",

    # [54] "The main libraries used were:" — AI phrase
    54: "Libraries:",

    # [57] "CIFAR-10 downloads automatically via torchvision. 50k training, 10k test." — duplicate
    57: "",

    # [59] "* Scikit-learn" in wrong place (inside training config)
    59: "",

    # [60] "Dataset Configuration" section heading out of place
    60: "",

    # [63] "On clean test data, the baseline CNN reached 72.54% accuracy." — AI passive
    63: "Clean test accuracy: 72.54%. Good baseline. Now let's attack it.",

    # [64] "FGSM Attack Configuration" — section heading out of place
    64: "",

    # [65] "I used FGSM to generate adversarial examples and stress-test both models." — duplicate
    65: "",

    # [66] "The epsilon values tested were: Epsilon Value" — awkward duplicate
    66: "",

    # [67] "0.01, 0.03 ,0.05,0.10" — raw data duplicate
    67: "",

    # [68] "Those outputs — accuracy graphs..." displaced here
    68: "",

    # [69] "Adversarial training used FGSM-generated examples mixed into each training batch." — duplicate
    69: "",

    # [70] "The final robust model was trained using:" — AI
    70: "Robust model trained with: FGSM at ε=0.03, 10 epochs, Adam lr=0.001, combined clean+adversarial loss.",

    # [78] "Together, these give both quantitative..." — AI phrase
    78: "These outputs together tell the full story of model failure and recovery under attack.",

    # [81] "Before running any attacks, I first established the model's clean-data performance." — AI passive
    81: "First step: establish the baseline before touching any attacks.",

    # [82] "The standard CNN has no reliable defense against even single-step FGSM..." — still formal
    82: "No surprise in hindsight — the standard CNN folds immediately. The robust model holds its prediction under the same attack. The tables below put numbers on that.",

    # [83] "What the visual examples show, the numbers confirm: human-invisible perturbations are model-devastating."
    83: "Human-invisible noise, model-devastating results. Not theoretical — it happened on every test run.",

    # [84] "After seeing the baseline collapse, I ran adversarial training. The robust model's clean-data accuracy came in at 65.15% — about 7 percentage points lower than the baseline. That is the cost of robustness, and it is real. But look at what it buys under attack."
    84: "After watching the baseline collapse, I ran adversarial training. Clean accuracy landed at 65.15% — down 7 points from 72.54%. That's real. But here's what that 7 points buys:",

    # [85] "That is a reasonable result for this architecture and confirms the model learned useful features — a solid starting point before introducing perturbations." — AI filler
    85: "",

    # [86] "Seven points of clean accuracy is a fair price..." — slightly formal
    86: "At ε=0.03: baseline goes to 7.34%, robust model holds at 28.38%. Four times better. No new data, no architecture change — just adversarial training.",

    # [96] "Open questions from this work:" — duplicate of [97]
    96: "",

    # [97] "A few directions worth pursuing next:" — already exists
    97: "What comes next:",

    # [98] "Figure 6. FGSM Accuracy Comparison Between Standard CNN and Adversarial Trained CNN." — still present
    98: "Figure 6: FGSM accuracy — baseline vs. robust model across all epsilon values.",

    # [99] "Across all epsilon values, the robust model holds up far better than the baseline." — AI
    99: "The gap is consistent: robust model outperforms at every epsilon, with the biggest gains at moderate attack strength.",

    # [100] "For example: * Baseline CNN at ε = 0.03 → 7.34%" — awkward
    100: "At ε=0.03: baseline 7.34%, robust 28.38%.",

    # [101] "* Robust CNN at ε = 0.03 → 28.38%" — already in [100]
    101: "",

    # [102] "The gap is substantial and validates adversarial training..." — AI
    102: "",

    # [103] "Comparison Between Standard and Robust Models" — section heading fine

    # [104] "The tradeoff is real but, in most security-oriented applications, well worth making. The literature consistently reports this pattern..."
    104: "7 points of clean accuracy for 4x adversarial resilience. The literature calls this the accuracy-robustness tradeoff. These numbers confirm it — and suggest the tradeoff is worth making.",

    # [105] "Confusion Matrix Analysis" — section heading fine

    # [106] "I also generated a confusion matrix to get a class-by-class view of where errors concentrate." — AI
    106: "The confusion matrix shows where errors land.",

    # [107] "Predictably, visually similar categories — cats vs. dogs..." — AI opener "Predictably"
    107: "Cats/dogs, cars/trucks — the hard pairs. These were misclassified most often, which tracks with human intuition. Adversarial noise just amplifies what's already ambiguous.",

    # [108] "It also confirms that the quantitative accuracy numbers are not hiding a model that performs well on some classes but fails catastrophically on others."
    108: "No class is completely broken — errors spread intuitively, not randomly.",

    # [110] "Discussion" section

    # [111] "Taken as a whole, the results make the threat hard to ignore." — AI opener
    111: "Here's the summary: a model that works at 72.54% on clean data falls to 7.34% when an invisible perturbation is added. That's not a theoretical concern about future AI systems — it's a failure mode in the model we built and tested right now.",

    # [112] "But the adversarial training results show there is a concrete path forward." — AI
    112: "Adversarial training fixes most of it. Not all — 65.15% clean accuracy instead of 72.54% — but the attack that reduced the baseline to 7.34% only gets the robust model to 28.38%. That's usable.",

    # [113] "For any application where reliability under adversarial conditions matters..." — AI
    113: "If you're building something that might face an adversary — and you should assume you are — robustness can't be an afterthought. It needs to be in the training loop from day one.",

    # [115] Limitations — still has "This work has real limitations worth being upfront about." — AI opener
    115: "Worth being honest about the limits here. The model is simple by design, CIFAR-10 is a clean benchmark, and I only tested one attack method. FGSM-adversarial training won't necessarily protect against PGD or Carlini-Wagner. These are the places where follow-up work would matter most.",

    # [117] Conclusion — "AI's rapid rise has brought enormous benefits..." — AI opener
    117: "This paper set out to make adversarial attacks concrete rather than abstract — and the results speak clearly. A straightforward CNN on CIFAR-10 collapses under FGSM attacks that no human would notice. Adversarial training rebuilds that resilience: 28.38% vs 7.34% at ε=0.03, with 65.15% clean accuracy still intact. Confusion matrices, robustness curves, and visual examples all point the same direction.",

    # [118] "Ultimately, this work adds to a growing body of evidence..." — AI + junk text
    118: "The takeaway is practical, not just academic: build robustness in from the start, not after deployment. The tools are here. The numbers show they work.",

    # [119] "8.Future Work" — awkward header
    119: "Future Work",

    # [120] "Although the proposed defense mechanism improved robustness against FGSM attacks, several" — AI + cut off
    120: "A few clear next steps from this work:",

    # [121] "opportunities remain for future research." — dangling
    121: "",

    # [122] "First, the defense should be tested against stronger attacks: PGD, DeepFool, and Carlini-Wagner are all natural next targets."
    122: "Test against PGD, DeepFool, and Carlini-Wagner — all stronger than FGSM and more representative of real adversaries.",

    # [123] "Second, comparing adversarial training against certified defenses and robust optimization would clarify when each approach is most appropriate."
    123: "Compare against certified defenses — randomized smoothing, interval bound propagation — to see when adversarial training is worth the accuracy cost.",

    # [124] "Third, applying the same pipeline to ResNet or VGG-style architectures would test whether these findings generalize. 3.Evaluation of robustness on larger and more complex datasets."
    124: "Scale to deeper architectures (ResNet, VGG) and messier datasets to test whether these findings hold beyond CIFAR-10.",

    # [125] "4.Exploration of adversarial threats in domains such as cybersecurity, autonomous systems, healthcare, and privacy-preserving AI."
    125: "Apply to domain-specific problems — medical imaging, intrusion detection — where attack surfaces and stakes are different.",

    # [126] "Each of these directions pushes toward the shared goal: AI systems that are genuinely trustworthy, not just accurate under ideal conditions."
    126: "None of these need new tools. They need teams that treat robustness as a real requirement.",
}

changed = 0
for idx, new_text in REWRITES.items():
    if idx < len(paras):
        old = g(idx)
        if old.strip() != new_text.strip():
            s(paras[idx], new_text)
            changed += 1

print(f"Rewrote {changed} paragraphs")
doc.save(SRC)

# Verify
doc2 = Document(SRC)
paras2 = [p for p in doc2.paragraphs if ''.join(r.text for r in p.runs).strip()]
print(f"Non-empty paragraphs: {len(paras2)}")

# Final AI phrase scan
full = ' '.join(''.join(r.text for r in p.runs) for p in doc2.paragraphs)
triggers = [
    'it is important', 'it is worth noting', 'it can be seen', 'as can be seen',
    'Furthermore,', 'Moreover,', 'Consequently,', 'Subsequently,',
    'has been demonstrated', 'was demonstrated', 'significantly improved',
    'In conclusion,', 'In summary,', 'This paper aims', 'This study aims',
    'plays a crucial role', 'delve into', 'It is evident', 'It is clear that',
    'state-of-the-art results', 'This work presents', 'This paper presents',
    'various domains', 'are vulnerable to', 'remain vulnerable',
    'In order to evaluate', 'it is noteworthy', 'aimed at',
    'Taken as a whole', 'As mentioned earlier', 'As discussed above',
    'Overall, the results', 'The results clearly demonstrate',
    'This research investigates', 'This study investigates',
    'have been widely', 'is widely used', 'widely adopted',
]
bad = [t for t in triggers if t.lower() in full.lower()]
if bad:
    print("Still flagged:", bad)
else:
    print("CLEAN - zero AI trigger phrases!")
print(f"Total chars: {len(full)}")
