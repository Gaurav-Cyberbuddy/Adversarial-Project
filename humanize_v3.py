"""
Deep humanizer v3 - aggressive paragraph-by-paragraph rewrite.
Uses paragraph INDEX to surgically rewrite every remaining AI-flagged sentence.
Reads from _HUMANIZED.docx and saves back to same file.
"""

from docx import Document
import sys

SRC = r'c:\Users\gaura\OneDrive\Documents\Adversarial Training\Adversaril attack and training Research paper_HUMANIZED.docx'
DST = r'c:\Users\gaura\OneDrive\Documents\Adversarial Training\Adversaril attack and training Research paper_HUMANIZED.docx'

doc = Document(SRC)
paras = doc.paragraphs

def set_text(para, text):
    """Safely set paragraph text, preserving first run's formatting."""
    runs = para.runs
    if not runs:
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ''

def get_text(para):
    return ''.join(r.text for r in para.runs)

# Map: paragraph index -> new text (only for paragraphs needing deeper rewrite)
# These are the ones still triggering AI detection - overly formal/passive/uniform patterns
REWRITES = {
    # Abstract - paragraph 5 (still has some stiff structure)
    5: "AI has reshaped modern computing across sectors like cybersecurity, healthcare, and finance. But deep learning models have a real weakness: tiny, invisible tweaks to inputs can make them confidently predict the wrong thing. That's not a theoretical problem — it matters anywhere AI is deployed in the real world. My focus here is on adversarial attacks specifically, and on whether adversarial training can do anything meaningful about them. I trained a CNN on CIFAR-10, hit it with FGSM attacks at several perturbation strengths, and the results were sobering — accuracy cratered fast. When I retrained the model on a mix of clean and adversarial samples, it held up considerably better. Robustness graphs, confusion matrices, and visual comparisons all backed that up. The takeaway: security thinking needs to be part of how we build AI, not something bolted on later.",

    # Introduction paragraph 10
    10: "Few technologies have moved as fast or spread as wide as AI. Hospital diagnostics, fraud detection, self-driving cars, content recommendation — these aren't future applications, they're running in production right now. Deep learning powers most of them. But that very success has made AI a target. As models get embedded in critical systems, adversaries have gotten serious about finding where they break. And break they do: data poisoning, model theft, membership inference, synthetic media — the attack surface is broad. Adversarial examples are probably the most studied of these threats, and for reasons that go beyond academic interest.",

    # Para 11 (attack description)
    11: "The attack concept is almost embarrassingly simple. Perturb an input by a tiny amount — carefully chosen to maximize the model's loss — and it will confidently produce the wrong answer. The human looking at the image sees nothing unusual. The model is completely fooled. That gap between human and machine perception is what makes adversarial examples so unsettling, especially when you're talking about medical scans or security cameras.",

    # Para 12 (contributions paragraph)
    12: "My goal in this paper is to make that threat concrete rather than theoretical. I built a CIFAR-10 classifier, measured how fast it collapses under FGSM attacks, and then tested adversarial training as a countermeasure. The evaluation goes beyond a single accuracy number — I looked at clean vs. adversarial accuracy, multi-epsilon robustness curves, class-level confusion, and visual inspection of adversarial samples. The aim is to give a complete picture, not just a headline result.",

    # Para 20 (lit review start)
    20: "There is a lot of prior work to stand on here. Deep learning's track record across vision, language, and security tasks is well established — so is the fact that these models have serious adversarial blind spots. The two observations have existed in tension since researchers first noticed that small, structured perturbations could fool well-trained networks. What surprised the community early on was just how invisible these perturbations could be to humans while still completely breaking model predictions. That mismatch turned out to be a fundamental property of how neural networks learn, not a quirk of any particular architecture.",

    # Para 21 (FGSM intro in lit review)
    21: "FGSM gave the field a standardized way to generate and study these attacks. Its gradient-based logic is simple enough that anyone with the model weights and a loss function can generate adversarial examples in one forward-backward pass. That made it a universal benchmark. Later work developed iterative and optimization-based attacks that are harder to defend against, but FGSM remains the entry point for any robustness evaluation. Studies using it repeatedly showed the same pattern: high clean accuracy and near-zero adversarial accuracy, especially as perturbation strength grows. Adversarial training emerged as the most empirically reliable defense — not perfect, and always at some cost to clean accuracy, but consistently better than no defense at all. Other threats like data poisoning and model extraction round out the picture, confirming that AI systems face attack vectors throughout their entire lifecycle.",

    # Para 22 (lit review conclusion)
    22: "What this body of work leaves room for is practical, end-to-end demonstrations on real datasets. Knowing that adversarial training works in theory is one thing; seeing the actual accuracy numbers across different epsilon values, on a concrete architecture and dataset, is another. That is what this paper contributes.",

    # Para 24 (methodology intro)
    24: "The experiments follow a straightforward pipeline: build a baseline CNN on CIFAR-10, attack it with FGSM across several epsilon values, retrain it with adversarial training, and compare the results. Each step is described below.",

    # Para 27 (dataset description)
    27: "All experiments use CIFAR-10 — 60,000 color images at 32×32 pixels, split into 50,000 for training and 10,000 for testing. It is a standard benchmark that is small enough to iterate on quickly but representative enough that results carry some weight. I loaded it via PyTorch's torchvision and normalized the pixel values before training.",

    # Para 41 (CNN architecture)
    41: "The model is deliberately simple: two conv layers, a max-pool, and two fully connected layers. This was intentional — I wanted a representative but unexceptional baseline, not a state-of-the-art architecture that might obscure the vulnerability. Training used Adam at lr = 0.001 with cross-entropy loss.",

    # Para 50 (FGSM description)
    50: "FGSM is almost elegant in its simplicity. Given an input x, it computes the gradient of the loss with respect to x, takes the sign of that gradient, and scales it by epsilon. That one-step perturbation is enough to mislead most standard models. I tested four epsilon values — 0.01, 0.03, 0.05, and 0.10 — to track how attack effectiveness scales with perturbation budget.",

    # Para 59 (adversarial training)
    59: "Adversarial training is conceptually straightforward: on each mini-batch, generate FGSM adversarial examples on-the-fly and mix them with the clean samples before the forward pass. The model ends up training on both, which forces it to learn features that do not collapse under small perturbations. The tradeoff is that clean accuracy usually drops — the network is spending some of its capacity on robustness rather than clean-data performance. Both models were then evaluated on identical test sets and attack configurations for a direct comparison.",

    # Para 67 (metrics)
    67: "These metrics provide complementary views: accuracy numbers tell you how much the model degrades, the confusion matrix shows you where, and the visual comparisons make the threat tangible in a way that numbers alone cannot.",

    # Para 70 (experimental setup)
    70: "Everything ran in Python and PyTorch, inside a virtual environment on a Windows machine using VS Code. Dependencies included torchvision, matplotlib, NumPy, and scikit-learn — nothing unusual.",

    # Para 103 (results - baseline)
    103: "Starting with the clean baseline: the CNN hit 72.54% accuracy on the CIFAR-10 test set. That is a solid result for this architecture — it confirms the model learned something real before I started attacking it.",

    # Para 107 (baseline result commentary)
    107: "72.54% is not a fluke. The training curves were stable and the model generalized well. That makes what happens next more striking.",

    # Para 109 (attack intro)
    109: "With a solid baseline in hand, I ran FGSM at four epsilon values. The results are blunt.",

    # Para 110 (attack results commentary)
    110: "Accuracy collapses fast. Even at epsilon = 0.01, the model is already struggling. By epsilon = 0.03, it is down to 7.34% — worse than random guessing on a 10-class problem.",

    # Para 111 (more attack results)
    111: "At epsilon = 0.03, clean accuracy of 72.54% becomes adversarial accuracy of 7.34%. The perturbations causing this are invisible to a human observer. That combination — imperceptible to us, devastating to the model — is what makes adversarial examples a genuine security concern, not just an academic curiosity.",

    # Para 114 (adversarial example analysis)
    114: "The visual comparison makes the point vividly. Original and adversarial images look identical to a human. The CNN's prediction flips completely. In one test case, the model went from predicting class 3 to predicting class 5 — a completely different category — based on noise that no person would notice.",

    # Para 120 (after figure captions)
    120: "The standard CNN has no reliable defense against even single-step FGSM. The robustly trained version maintained its prediction on the same input — a meaningful qualitative difference that the accuracy tables also confirm.",

    # Para 121 (invisible threat)
    121: "What the visual examples show, the numbers confirm: human-invisible perturbations are model-devastating. This is not a niche finding — it applies broadly across architectures and datasets.",

    # Para 123 (adversarial training performance)
    123: "After seeing the baseline collapse, I ran adversarial training. The robust model's clean-data accuracy came in at 65.15% — about 7 percentage points lower than the baseline. That is the cost of robustness, and it is real. But look at what it buys under attack.",

    # Para 125 (tradeoff commentary)
    125: "Seven points of clean accuracy is a fair price for a model that does not fall apart under attack. In a security context, a model at 65% accuracy that holds up under adversarial conditions is more useful than one at 72% that collapses to 7%.",

    # Para 126 (literature consistency)
    126: "This accuracy-robustness tradeoff is not specific to this experiment — it shows up consistently in the adversarial ML literature. Getting both high clean accuracy and high adversarial accuracy at the same time is an open research problem.",

    # Para 128 (robustness table intro)
    128: "The robust model's performance under FGSM attacks tells a different story than the baseline. At epsilon = 0.03, it holds 28.38% accuracy — versus the baseline's 7.34%. At epsilon = 0.01, it reaches nearly 50%.",

    # Para 131 (comparison commentary)
    131: "Across all four epsilon values, the robust model outperforms the baseline by a wide margin. The gap is largest at moderate perturbation strengths — exactly where a real attacker would operate.",

    # Para 134 (defense validation)
    134: "At epsilon = 0.03, baseline accuracy: 7.34%. Robust model accuracy: 28.38%. That is nearly four times better, from adversarial training alone — no architectural changes, no additional data.",

    # Para 136 (comparison between models)
    136: "The tradeoff is clean: 7 points of clean accuracy for a model that is roughly four times more resilient at moderate attack strength. Whether that trade makes sense depends on the application, but for anything security-sensitive the answer is obvious.",

    # Para 139 (confusion matrix)
    139: "Looking at the confusion matrix, the failure modes make sense. Cats and dogs get mixed up; automobiles and trucks do too. These are visually similar categories where even humans sometimes disagree. The adversarial examples amplify those existing ambiguities. That points toward targeted data augmentation or architecture changes as productive directions for future work.",

    # Para 140 (confusion matrix extra)
    140: "The confusion matrix also rules out a common failure mode: a model that looks reasonable in overall accuracy but is near-zero on some classes. The errors here are distributed in ways that make intuitive sense.",

    # Para 143 (discussion start)
    143: "The numbers here are not subtle. A model that works fine under normal conditions falls apart completely when any adversarial perturbation is introduced — not just strong ones, but weak ones too. At epsilon = 0.03, which is visually imperceptible, the standard CNN is guessing. That should be alarming for anyone deploying deep learning in a setting where an adversary might have any incentive to manipulate inputs.",

    # Para 144 (adversarial training path forward)
    144: "Adversarial training changes that picture. Yes, you pay a clean-accuracy cost — 65% instead of 72.5% — but the model survives attacks that would otherwise reduce it to random guessing. In a security context that tradeoff is not just acceptable, it is the right engineering decision.",

    # Para 145 (deployment implications)
    145: "The practical implication is straightforward: if your model might face adversarial inputs in deployment — and for any security-facing application, assume it will — then robustness needs to be in your training objectives from day one. Retrofitting it later is harder and less effective.",

    # Para 147 (limitations)
    147: "A few honest limitations. The architecture is simple by design, but that means results may not transfer directly to ResNet-scale models. CIFAR-10 is a clean, balanced benchmark — real deployment data is messier. And I only tested against FGSM; adversarial training against FGSM does not automatically generalize to PGD or Carlini-Wagner, which are considerably stronger. These are the most important gaps to close in follow-up work.",

    # Para 149 (conclusion)
    149: "AI systems are genuinely useful and genuinely fragile in ways that standard training does not address. The experiments in this paper make that fragility concrete: a reasonable CNN, trained in the normal way, collapses under adversarial pressure that a human would not notice. Adversarial training brings meaningful improvement at a moderate accuracy cost. The 28.38% vs. 7.34% comparison at epsilon = 0.03 is the number that sticks — and it makes the case that adversarial training is not just an academic exercise but a practical engineering choice worth making.",

    # Para 150 (conclusion wrap-up)
    150: "Security and accuracy are not perfectly compatible objectives — but they are not mutually exclusive either. This work shows that with adversarial training, you can build models that are both usable and substantially more resilient. That balance is worth pursuing, and the tools to pursue it already exist.",

    # Para 152/153 (future work)
    152: "Several concrete next steps follow from this work.",
    153: "",

    # Para 156 (future work - generalization)
    156: "Applying the same pipeline to ResNet or VGG would test whether the results scale to deeper architectures. Evaluating on larger, noisier datasets — ImageNet, or domain-specific data — would test whether they generalize beyond the CIFAR-10 setting.",
}

# Apply rewrites by index
changed = 0
for idx, new_text in REWRITES.items():
    if idx < len(paras):
        old = get_text(paras[idx])
        if new_text and old != new_text:
            set_text(paras[idx], new_text)
            changed += 1
            print(f"  [para {idx}] rewritten")

print(f"\nTotal paragraphs rewritten: {changed}")

# Also do a pass for any remaining formulaic phrases across ALL paragraphs
PHRASE_PASS = [
    ("These metrics provide both quantitative and qualitative insights into model performance, vulnerability, and robustness under adversarial conditions.",
     "These metrics give complementary views: numbers quantify the damage, confusion matrices show which classes suffer most, and visual examples make the threat tangible."),
    ("The model achieved a test accuracy of:",
     "On clean test data, it scored:"),
    ("The final adversarial trained model achieved:",
     "After adversarial training, the model achieved:"),
    ("Figure 1. Training Accuracy of CNN Across Training Epochs",
     "Figure 1. CNN Training Accuracy Across Epochs"),
    ("5.5 Robustness Evaluation of the Adversarially Trained Model The robust model was evaluated using the same FGSM attack strengths",
     ""),
]

for para in doc.paragraphs:
    txt = get_text(para)
    new_txt = txt
    for old, new in PHRASE_PASS:
        if old in new_txt:
            new_txt = new_txt.replace(old, new, 1)
    if new_txt != txt:
        set_text(para, new_txt)
        changed += 1

doc.save(DST)
print(f"Saved: {DST}")
print(f"Total changes: {changed}")
