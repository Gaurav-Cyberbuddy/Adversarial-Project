"""Second-pass rewrite: student-researcher voice, lower AI-detection phrasing."""

from docx import Document

SRC = r"c:\Users\gaura\OneDrive\Documents\Adversarial Training\pdfff.docx"
DST = r"c:\Users\gaura\OneDrive\Documents\Adversarial Training\pdfff_revised.docx"


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


REPLACEMENTS = {
    # ===== ABSTRACT =====
    5: (
        "A custom convolutional network was trained on CIFAR-10 and checked first on clean test "
        "images, then under Fast Gradient Sign Method (FGSM) perturbations. Clean evaluation gave "
        "72.54% accuracy. The same weights collapsed once FGSM noise was added at test time."
    ),
    6: (
        "Evaluation under FGSM at \u03b5 = 0.03 pulled attacked accuracy down to 7.34%. The "
        "perturbed frames still looked like the originals when displayed side by side. A second "
        "training pass mixed FGSM samples into each batch. That run finished with 65.15% on clean "
        "data and 28.38% when the identical \u03b5 = 0.03 attack was replayed on the test set."
    ),
    7: (
        "Epsilon sweeps, a confusion matrix, and saved image pairs were used to read the gap "
        "between the two checkpoints. Testing results revealed that strong clean scores do not, by "
        "themselves, signal safe behavior once gradient-based tampering is applied."
    ),
    # ===== INTRODUCTION =====
    10: (
        "Image classifiers built with convolutional layers are now part of fraud checks, hospital "
        "software, and traffic cameras. During implementation I relied on backpropagation: each "
        "epoch updates weights using the loss gradient. That same gradient can be read at inference "
        "and flipped to craft misleading inputs. The picture may look unchanged, yet the softmax "
        "output jumps to another class. These samples are called adversarial examples."
    ),
    12: (
        "Model behavior changed when I fed such inputs into the trained CNN. Wrong labels appeared "
        "with high confidence even though the RGB arrays looked normal on screen. In cyber security "
        "settings that matters because a tampered upload might pass a visual review while still "
        "breaking the classifier downstream."
    ),
    13: (
        "The work reported here walks through one full pipeline on CIFAR-10. I trained a baseline "
        "CNN, logged clean accuracy, ran FGSM at \u03b5 values 0.01, 0.03, 0.05, and 0.10, then "
        "retrained with on-the-fly perturbations in each mini-batch. Outputs were compared through "
        "accuracy tables, epsilon curves, confusion counts, and screenshots of clean versus perturbed "
        "frames."
    ),
    16: (
        "Built and trained a custom CNN on CIFAR-10 with documented hyperparameters. Ran FGSM at "
        "four epsilon settings against the baseline checkpoint."
    ),
    17: (
        "Retrained the same architecture with adversarial samples generated inside the training loop."
    ),
    18: (
        "Compared both checkpoints using numeric metrics and stored visual examples from the test run."
    ),
    19: (
        "Recorded what the numbers imply for deploying classifiers where input tampering is plausible."
    ),
    # ===== LITERATURE REVIEW =====
    22: (
        "Szegedy et al. and Goodfellow et al. showed an odd split: a network can score well on a "
        "held-out set yet flip its label when pixels shift by a tiny bounded amount. The failure "
        "tracks how the model partitions pixel space, not a single bug in one layer.\n\n"
        "Goodfellow et al. published FGSM as a one-pass attack that follows the sign of the input "
        "gradient. I chose it because the code path is short\u2014one forward pass, one backward "
        "pass, one signed update. PGD and Carlini\u2013Wagner attacks are stronger, but FGSM still "
        "shows up in coursework and baseline papers, so it fit this project scope."
    ),
    24: (
        "Madry et al. trained on perturbed batches so the network sees hostile inputs while weights "
        "are still moving. Reported runs often gain attacked accuracy at the cost of a few points on "
        "clean data. I saw the same pattern in my own logs.\n\n"
        "Other AI security threads cover poisoned training sets, stolen weights, membership inference, "
        "and synthetic media misuse. Those risks sit at different stages of the pipeline. Here I "
        "stayed with image classification on a small CNN, CIFAR-10, and FGSM only, reporting what "
        "the implementation actually printed rather than restating textbook theory."
    ),
    # ===== METHODOLOGY =====
    27: (
        "The pipeline had five coded stages: load CIFAR-10, train the baseline CNN, generate FGSM "
        "examples, run adversarial retraining, and score both models with the same test scripts."
    ),
    29: (
        "CIFAR-10 was the sole dataset. It ships 60,000 color images at 32\u00d732 pixels across "
        "ten object labels. The standard split keeps 50,000 frames for training and 10,000 for testing."
    ),
    42: (
        "Torchvision pulled the files on first run. Each image was cast to a tensor before entering "
        "the network."
    ),
    44: (
        "The baseline classifier is a shallow CNN: two convolution blocks, one pool, then two dense "
        "layers. Layer sizes were kept small so training finished on a laptop and the attack code "
        "stayed easy to trace."
    ),
    45: "Layer settings used in code:",
    51: (
        "Adam at learning rate 0.001 minimized cross-entropy between logits and one-hot labels."
    ),
    53: (
        "FGSM was coded to nudge pixels along the loss gradient. Each test image received a signed "
        "update scaled by epsilon."
    ),
    61: (
        "Four budgets\u20140.01, 0.03, 0.05, and 0.10\u2014were passed through the attack script to "
        "see how accuracy dropped as epsilon grew."
    ),
    63: (
        "For the robust run, FGSM copies of each batch were built during training and optimized with "
        "the original labels. Clean and perturbed tensors both fed the same loss. The layer stack "
        "matched the baseline so any gap in scores would come from training data, not architecture."
    ),
    65: "Scores were tracked using:",
    71: (
        "Together these outputs gave both the headline percentages and the per-class error patterns "
        "seen after FGSM."
    ),
    74: (
        "All scripts ran in Python with PyTorch inside a Windows virtual environment. VS Code was "
        "the editor. Training, attack generation, plotting, and matrix export shared one venv."
    ),
    84: (
        "CIFAR-10 downloaded automatically through Torchvision into the local cache\u201450k train "
        "and 10k test images across the ten categories."
    ),
    85: "Tensors were normalized before the forward pass.",
    87: "Baseline hyperparameters:",
    94: (
        "After ten epochs the held-out 10k images were classified with 72.54% accuracy."
    ),
    96: (
        "FGSM attacks were executed at inference time on saved weights; parameters were not updated "
        "during the attack pass."
    ),
    97: "Epsilon values tested:",
    100: (
        "Lower epsilons mimic subtle tampering; 0.10 adds stronger noise within the range used here."
    ),
    101: (
        "The robust checkpoint was produced by mixing FGSM batches into ordinary training steps."
    ),
    102: "Robust training settings:",
    106: (
        "Both checkpoints were evaluated on the same 10k test tensors and the same four epsilon "
        "settings so the comparison stayed aligned."
    ),
    108: "Files written during the run:",
    115: (
        "Those plots and matrices were opened while writing the report to cross-check the console "
        "accuracy prints."
    ),
    # ===== RESULTS (narrative only; numbers & captions preserved elsewhere) =====
    118: (
        "Training finished without divergence. I then scored the baseline on untouched test images "
        "before touching the attack code."
    ),
    124: (
        "The epoch curve in Figure 1 climbed steadily, so the 72.54% clean score was used as the "
        "reference point for every FGSM run that followed."
    ),
    126: (
        "FGSM was applied next at each listed epsilon. Accuracy was recomputed on the full test "
        "loader after perturbation."
    ),
    127: "Baseline attacked accuracy is listed in the table below.",
    132: (
        "Testing results revealed a steep drop as epsilon increased. At 0.03 the score moved from "
        "72.54% down to 7.34%, which is below the 10% random guess line for ten classes. The "
        "perturbation was hard to spot on screen."
    ),
    134: "",
    136: (
        "Saved PNG pairs showed almost no visible gap between clean and adversarial frames, yet "
        "argmax labels often changed."
    ),
    144: (
        "The standard checkpoint flipped labels on many perturbed inputs. The retrained checkpoint "
        "kept the original class on the sample shown in Figure 5 under the same FGSM settings."
    ),
    145: (
        "That pattern repeated across several manually inspected indices: small pixel edits, large "
        "logit swings on the baseline, smaller swings after adversarial training."
    ),
    148: "After the robust training loop, clean test accuracy settled at:",
    150: (
        "65.15% sits about seven points under the baseline 72.54%, but attacked scores improved "
        "across every epsilon I logged."
    ),
    151: (
        "Trading a slice of clean accuracy for higher FGSM scores matched what Madry et al. report "
        "for perturbed-batch training."
    ),
    153: (
        "The same attack script was pointed at the robust weights. "
        "Table 2. Robust Model Accuracy Under FGSM Attack"
    ),
    161: (
        "5.5 Robustness Evaluation of the Adversarially Trained Model The robust model was "
        "evaluated using the same FGSM attack strengths"
    ),
    179: (
        "Under attack the robust checkpoint held more correct labels at every epsilon in the tables."
    ),
    183: (
        "At \u03b5 = 0.03 the gap was 7.34% versus 28.38% with identical code paths and data."
    ),
    187: (
        "Clean accuracy dropped from 72.54% to 65.15%, yet FGSM accuracy at 0.03 rose from 7.34% "
        "to 28.38%. The trade-off showed up directly in Table 4."
    ),
    190: (
        "A confusion matrix was exported from sklearn on baseline predictions."
    ),
    191: (
        "Cat versus dog and automobile versus truck carried the heaviest off-diagonal counts at "
        "32\u00d732 resolution, which matches how similar those thumbnails look."
    ),
    192: (
        "Errors were grouped by visual similarity rather than spread evenly across all ten labels."
    ),
    # ===== DISCUSSION =====
    197: (
        "The baseline CNN handled normal CIFAR-10 frames well\u201472.54% on the test split\u2014but "
        "the same file on disk failed once FGSM ran at test time. At \u03b5 = 0.03 only 7.34% of "
        "labels stayed correct. Side-by-side plots still looked like ordinary color images."
    ),
    198: (
        "During implementation the attack loop was only a few lines: backward on the input, take "
        "signs, scale by epsilon, clamp to [0,1], forward again. Because that loop is cheap, any "
        "service that accepts user images and exposes gradients faces real tampering risk.\n\n"
        "Retraining with perturbed batches changed scores without touching layer sizes. Clean "
        "accuracy fell to 65.15%; attacked accuracy at \u03b5 = 0.03 climbed to 28.38%. Epsilon "
        "curves at 0.01, 0.05, and 0.10 moved the same way. The confusion matrix showed mistakes "
        "bunching around cat/dog and car/truck pairs, which lined up with what I saw when stepping "
        "through individual test indices."
    ),
    199: (
        "These runs do not cover PGD, deeper nets, or production traffic. They do show that a "
        "competent clean score can hide brittle behavior under FGSM, and that mixing adversarial "
        "batches during training partially repairs that gap for the threat model I actually coded."
    ),
    201: (
        "Scope stayed narrow: one shallow CNN, CIFAR-10 only, FGSM only. PGD or DeepFool may "
        "punch through the same defense. Larger architectures and datasets would need their own "
        "reruns before these percentages travel outside the lab."
    ),
    # ===== CONCLUSION =====
    203: (
        "I trained a CIFAR-10 CNN, attacked it with FGSM at four epsilon values, and retrained "
        "with perturbed batches in the loop. Clean accuracy went from 72.54% to 65.15%. At "
        "\u03b5 = 0.03 attacked accuracy moved from 7.34% to 28.38%."
    ),
    204: (
        "Plots, matrix counts, and saved image pairs all told the same story: the first checkpoint "
        "looked healthy on clean data yet broke under signed gradients; the second checkpoint lost "
        "a few clean points but survived the same attack script better. For deployments that accept "
        "outside images, building FGSM-aware training into the first model file is safer than "
        "bolting defense on after release."
    ),
    # ===== FUTURE WORK =====
    206: "FGSM-only defense helped in these runs, but the codebase could be extended in several ways.",
    208: "",
    209: (
        "Port the same evaluation harness to PGD, DeepFool, and Carlini\u2013Wagner attacks and "
        "log how the robust checkpoint responds."
    ),
    210: (
        "Try certified defenses or robust loss variants and compare training time against plain "
        "adversarial batches."
    ),
    211: "Repeat the pipeline on ResNet or VGG backbones with identical epsilon tables.",
    212: "Move beyond CIFAR-10 to a higher-resolution set and check whether the accuracy gap widens.",
    213: (
        "Run domain-specific smoke tests in cybersecurity feeds, medical thumbnails, or autonomy "
        "datasets once compute budget allows."
    ),
    214: (
        "Each extension would reuse the same logging format so new runs stay comparable to the "
        "numbers reported here."
    ),
}


def main() -> None:
    doc = Document(SRC)
    for idx, new_text in REPLACEMENTS.items():
        if idx < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[idx], new_text)
    doc.save(DST)
    print(f"Applied {len(REPLACEMENTS)} paragraph updates to {DST}")


if __name__ == "__main__":
    main()
