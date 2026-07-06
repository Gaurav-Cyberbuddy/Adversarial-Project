"""Fix IEEE-style formatting in pdfff_revised.docx."""

from __future__ import annotations

from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC = r"c:\Users\gaura\OneDrive\Documents\Adversarial Training\pdfff_revised.docx"
DST = r"c:\Users\gaura\OneDrive\Documents\Adversarial Training\pdfff_revised.docx"
FALLBACK = r"c:\Users\gaura\OneDrive\Documents\Adversarial Training\pdfff_formatted.docx"


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_style(paragraph: Paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        paragraph.style = "Normal"


def set_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def clean_table_cells(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = p.text.strip()


def apply_heading_map(doc: Document) -> None:
    h1 = {
        "ABSTRACT",
        "Introduction",
        "Literature Review",
        "Methodology",
        "Results and Discussion",
        "Discussion",
        "Limitations of the Study",
        "Conclusion",
        "8. Future Work",
        "9. References",
    }
    h2 = {
        "Dataset",
        "Convolutional Neural Network Architecture",
        "Fast Gradient Sign Method (FGSM)",
        "Adversarial Training",
        "Evaluation Metrics",
        "Experimental Setup",
        "Generated Artifacts",
        "Baseline CNN Performance",
        "Impact of FGSM Adversarial Attacks",
        "Adversarial Example Analysis",
        "Adversarial Training Performance",
        "Robustness Evaluation of the Adversarial Trained Model",
        "Comparison Between Standard and Robust Models",
        "Confusion Matrix Analysis",
    }
    h3 = {
        "Development Environment",
        "Dataset Configuration",
        "CNN Training Configuration",
        "FGSM Attack Configuration",
    }

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text in h1:
            set_style(p, "Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text in h2:
            set_style(p, "Heading 2")
        elif text in h3:
            set_style(p, "Heading 3")


def apply_text_fixes(doc: Document) -> None:
    fixes = {
        "8.Future Work": "8. Future Work",
        "9.References": "9. References",
        "0.01, 0.03 ,0.05,0.10": "0.01, 0.03, 0.05, 0.10",
        "65.15% Classification Accuracy": "65.15%",
        "The model achieved a test accuracy of:": "The model achieved a test accuracy of 72.54%.",
        "Baseline attacked accuracy is listed in the table below.": (
            "Table 2 lists baseline accuracy under FGSM attack."
        ),
        "The same attack script was pointed at the robust weights. "
        "Table 2. Robust Model Accuracy Under FGSM Attack": (
            "Table 3 lists robust-model accuracy under the same FGSM attack strengths."
        ),
        "5.5 Robustness Evaluation of the Adversarially Trained Model "
        "The robust model was evaluated using the same FGSM attack strengths": "",
        "For example: * Baseline CNN at ε = 0.03 → 7.34%": (
            "At ε = 0.03, baseline CNN accuracy was 7.34%."
        ),
        "Built and trained a custom CNN on CIFAR-10 with documented hyperparameters. "
        "Ran FGSM at four epsilon settings against the baseline checkpoint.": (
            "Built and trained a custom CNN on CIFAR-10; ran FGSM at four epsilon "
            "settings against the baseline checkpoint."
        ),
    }

    for p in doc.paragraphs:
        text = p.text.strip()
        if text in fixes:
            set_text(p, fixes[text])


def normalize_lists(doc: Document) -> None:
    list_prefixes = ("* ", "• ")
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith(list_prefixes):
            set_style(p, "List Paragraph")
            if text.startswith("* "):
                set_text(p, text[2:].strip())


def remove_redundant_accuracy_block(doc: Document) -> None:
    """Collapse standalone 72.54% line after merged accuracy sentence."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "The model achieved a test accuracy of 72.54%.":
            if i + 1 < len(doc.paragraphs) and doc.paragraphs[i + 1].text.strip() == "72.54%":
                delete_paragraph(doc.paragraphs[i + 1])
            break


def remove_empty_paragraphs(doc: Document) -> int:
    removed = 0
    for p in reversed(doc.paragraphs):
        if not p.text.strip():
            delete_paragraph(p)
            removed += 1
    return removed


def add_table_captions(doc: Document) -> None:
    """Insert IEEE-style table captions immediately before each table."""
    captions = [
        "TABLE I\nBASELINE CNN TRAINING CONFIGURATION",
        "TABLE II\nADVERSARIAL TRAINING CONFIGURATION",
        "TABLE III\nBASELINE CNN ACCURACY UNDER FGSM ATTACK",
        "TABLE IV\nROBUST MODEL ACCURACY UNDER FGSM ATTACK",
        "TABLE V\nCOMPARISON BETWEEN STANDARD AND ROBUST MODELS",
    ]

    body = doc.element.body
    tables = body.findall(qn("w:tbl"))
    for table_el, caption in zip(tables, captions):
        prev = table_el.getprevious()
        if prev is not None and prev.tag == qn("w:p"):
            para = Paragraph(prev, doc)
            if para.text.strip().upper().startswith("TABLE "):
                set_text(para, caption)
                set_style(para, "Caption")
                continue

        new_p = deepcopy(doc.paragraphs[0]._element)
        for child in list(new_p):
            if child.tag == qn("w:r"):
                new_p.remove(child)
        table_el.addprevious(new_p)
        para = Paragraph(new_p, doc)
        set_text(para, caption)
        try:
            set_style(para, "Caption")
        except Exception:
            set_style(para, "Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def format_figure_captions(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("Figure "):
            try:
                set_style(p, "Caption")
            except Exception:
                set_style(p, "Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def format_contributions(doc: Document) -> None:
    in_contributions = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "The primary contributions of this work are as follows:":
            in_contributions = True
            set_style(p, "Body Text")
            continue
        if in_contributions:
            if text.startswith("Built and trained") or text.startswith("Retrained") or (
                text.startswith("Compared both")
            ) or text.startswith("Recorded what"):
                set_style(p, "List Paragraph")
            else:
                break


def format_future_work_list(doc: Document) -> None:
    in_future = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "8. Future Work":
            in_future = True
            continue
        if in_future and text == "9. References":
            break
        if in_future and text and not text.startswith("FGSM-only") and not text.startswith(
            "Each extension"
        ):
            set_style(p, "List Paragraph")


def format_title_block(doc: Document) -> None:
    if len(doc.paragraphs) >= 3:
        set_style(doc.paragraphs[0], "Title")
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_style(doc.paragraphs[1], "Normal")
        doc.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_style(doc.paragraphs[2], "Normal")
        doc.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.CENTER


def format_references(doc: Document) -> None:
    in_refs = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "9. References":
            in_refs = True
            continue
        if in_refs and text:
            set_style(p, "List Paragraph")


def main() -> None:
    doc = Document(SRC)

    format_title_block(doc)
    apply_text_fixes(doc)
    normalize_lists(doc)
    remove_redundant_accuracy_block(doc)
    format_contributions(doc)
    format_future_work_list(doc)
    format_references(doc)
    format_figure_captions(doc)
    apply_heading_map(doc)
    clean_table_cells(doc)

    removed = remove_empty_paragraphs(doc)
    add_table_captions(doc)

    try:
        doc.save(DST)
        out = DST
    except PermissionError:
        doc.save(FALLBACK)
        out = FALLBACK

    print(f"Saved formatted document to {out}")
    print(f"Removed {removed} empty paragraphs")
    print(f"Paragraphs remaining: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
