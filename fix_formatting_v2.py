"""Second formatting pass: headings, merges, section order."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

SRC = r"c:\Users\gaura\OneDrive\Documents\Adversarial Training\pdfff_revised.docx"
DST = SRC
FALLBACK = r"c:\Users\gaura\OneDrive\Documents\Adversarial Training\pdfff_formatted.docx"

MAIN_HEADINGS = {
    "ABSTRACT",
    "Introduction",
    "Literature Review",
    "Methodology",
    "Results and Discussion",
    "Discussion",
    "Limitations of the Study",
    "Conclusion",
    "8. Future Work",
    "9. References",
}

SUB_HEADINGS = {
    "Dataset",
    "Convolutional Neural Network Architecture",
    "Fast Gradient Sign Method (FGSM)",
    "Adversarial Training",
    "Evaluation Metrics",
    "Experimental Setup",
    "Generated Artifacts",
    "Baseline CNN Performance",
    "Impact of FGSM Adversarial Attacks",
    "Adversarial Example Analysis",
    "Adversarial Training Performance",
    "Robustness Evaluation of the Adversarial Trained Model",
    "Comparison Between Standard and Robust Models",
    "Confusion Matrix Analysis",
    "Development Environment",
    "Dataset Configuration",
    "CNN Training Configuration",
    "FGSM Attack Configuration",
}


def set_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _bold_paragraph(paragraph: Paragraph, size_pt: int) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text = paragraph.text
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)


def style_main_heading(paragraph: Paragraph) -> None:
    try:
        paragraph.style = "Heading 1"
    except KeyError:
        _bold_paragraph(paragraph, 14)


def style_sub_heading(paragraph: Paragraph) -> None:
    _bold_paragraph(paragraph, 12)


def style_caption(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run(paragraph.text)
    run.bold = True
    run.font.size = Pt(10)


def find_paragraph(doc: Document, text: str) -> Paragraph | None:
    target = text.strip()
    for p in doc.paragraphs:
        if p.text.strip() == target:
            return p
    return None


def move_paragraph_before(doc: Document, moving: Paragraph, anchor: Paragraph) -> None:
    body = doc.element.body
    moving_el = moving._element
    anchor_el = anchor._element
    body.remove(moving_el)
    anchor_el.addprevious(moving_el)


def apply_heading_styles(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text in MAIN_HEADINGS:
            style_main_heading(p)
        elif text in SUB_HEADINGS:
            style_sub_heading(p)
        elif text.startswith("Figure "):
            style_caption(p)
        elif text.startswith("TABLE "):
            style_caption(p)


def apply_content_fixes(doc: Document) -> None:
    replacements = {
        "After the robust training loop, clean test accuracy settled at:": (
            "After the robust training loop, clean test accuracy settled at 65.15%."
        ),
        "65.15%": "",  # merged above; clear duplicate line
        "FGSM was applied next at each listed epsilon. Accuracy was recomputed on the full test loader after perturbation.": (
            "FGSM was applied next at each listed epsilon. Accuracy was recomputed on the full "
            "test loader after perturbation. Table III summarizes baseline accuracy under FGSM."
        ),
        "Robustness Evaluation of the Adversarial Trained Model": (
            "Robustness Evaluation of the Adversarial Trained Model"
        ),
        "Clean accuracy dropped from 72.54% to 65.15%, yet FGSM accuracy at 0.03 rose from 7.34% to 28.38%. The trade-off showed up directly in Table 4.": (
            "Clean accuracy dropped from 72.54% to 65.15%, yet FGSM accuracy at 0.03 rose from "
            "7.34% to 28.38%. The trade-off showed up directly in Table V."
        ),
        "At ε = 0.03, baseline CNN accuracy was 7.34%.": (
            "At ε = 0.03, baseline CNN accuracy was 7.34% and robust CNN accuracy was 28.38%."
        ),
        "Robust CNN at ε = 0.03 → 28.38%": "",
        "At ε = 0.03 the gap was 7.34% versus 28.38% with identical code paths and data.": "",
    }

    for p in doc.paragraphs:
        text = p.text.strip()
        if text in replacements:
            set_text(p, replacements[text])
            if not replacements[text]:
                p._element.getparent().remove(p._element)


def insert_before(doc: Document, anchor: Paragraph, text: str) -> Paragraph:
    new_p = anchor._element.makeelement(qn("w:p"), {})
    anchor._element.addprevious(new_p)
    para = Paragraph(new_p, doc)
    para.add_run(text)
    return para


def add_robust_table_intro(doc: Document) -> None:
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("TABLE IV"):
            anchor = p
            break
    if anchor is None:
        return
    prev = anchor._element.getprevious()
    if prev is not None:
        prev_para = Paragraph(prev, doc)
        if "same attack script" in prev_para.text:
            return
    insert_before(
        doc,
        anchor,
        "The same FGSM attack script was run on the robust checkpoint. Results are listed below.",
    )


def reorder_comparison_section(doc: Document) -> None:
    heading = find_paragraph(doc, "Comparison Between Standard and Robust Models")
    table_cap = find_paragraph(doc, "TABLE V\nCOMPARISON BETWEEN STANDARD AND ROBUST MODELS")
    if heading and table_cap:
        move_paragraph_before(doc, heading, table_cap)


def remove_empty_paragraphs(doc: Document) -> None:
    for p in reversed(doc.paragraphs):
        if not p.text.strip():
            p._element.getparent().remove(p._element)


def main() -> None:
    doc = Document(SRC)
    apply_content_fixes(doc)
    remove_empty_paragraphs(doc)
    add_robust_table_intro(doc)
    reorder_comparison_section(doc)
    apply_heading_styles(doc)

    try:
        doc.save(DST)
        out = DST
    except PermissionError:
        doc.save(FALLBACK)
        out = FALLBACK

    print(f"Formatting v2 saved to {out}")


if __name__ == "__main__":
    main()
