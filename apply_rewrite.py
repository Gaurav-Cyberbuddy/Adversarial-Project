"""Apply rewritten narrative sections to pdfff.docx."""

from docx import Document
from copy import deepcopy

SRC = r"c:\Users\gaura\OneDrive\Documents\Adversarial Training\pdfff.docx"
DST = r"c:\Users\gaura\OneDrive\Documents\Adversarial Training\pdfff.docx"

REPLACEMENTS = {
    # --- ABSTRACT ---
    5: (
        "Image classifiers built with convolutional neural networks often report strong accuracy "
        "on benchmark test sets, yet that score can collapse when inputs receive bounded "
        "pixel-level perturbations. In this project, a compact CNN was trained on CIFAR-10 and "
        "evaluated under the Fast Gradient Sign Method (FGSM) before and after adversarial training."
    ),
    6: (
        "On clean test images, the baseline model reached 72.54% accuracy. FGSM at \u03b5 = 0.03 "
        "reduced attacked accuracy to 7.34%, even though the perturbed images remained visually "
        "close to the originals. A second model was then trained with FGSM-generated samples mixed "
        "into each batch. That robust variant achieved 65.15% clean accuracy and 28.38% accuracy "
        "under the same \u03b5 = 0.03 attack\u2014substantially higher than the baseline under "
        "identical conditions."
    ),
    7: (
        "Robustness curves across four epsilon settings, confusion matrix analysis, and side-by-side "
        "visual comparisons of clean versus perturbed inputs were used to interpret the numeric gap "
        "between the two models. The experiments show that robustness should be treated as a "
        "training-time design choice rather than assumed from high clean accuracy alone."
    ),
    8: "",  # merged into paragraphs above
    # --- INTRODUCTION ---
    10: (
        "Convolutional neural networks are now embedded in software used for medical imaging review, "
        "payment fraud screening, traffic monitoring, and intrusion detection. During "
        "implementation, these models rely on gradient-based optimization: each training step "
        "adjusts weights according to how sensitive the loss is to small input changes. That same "
        "gradient signal can be inverted at inference time. An attacker who can supply or query "
        "inputs may push an image along the loss gradient so the network outputs a different class "
        "while the picture still appears normal to a human observer."
    ),
    12: (
        "Such manipulated inputs are called adversarial examples. They are not random noise "
        "artifacts; they are deliberately constructed failures. In a security-sensitive pipeline, "
        "a wrong label may pass unnoticed because the input still looks legitimate. That property "
        "makes adversarial risk relevant to cyber security work, where model outputs may gate "
        "access, trigger alerts, or influence automated decisions."
    ),
    13: (
        "This paper documents a controlled experiment on CIFAR-10 image classification. A baseline "
        "CNN was trained and scored on clean test data. FGSM attacks were then executed at "
        "perturbation strengths \u03b5 \u2208 {0.01, 0.03, 0.05, 0.10}. After measuring the drop in "
        "attacked accuracy, a second training run incorporated on-the-fly FGSM perturbations into "
        "each mini-batch. Both models were compared using accuracy tables, multi-epsilon robustness "
        "plots, confusion matrices, and stored image pairs showing prediction changes."
    ),
    14: "",  # content moved to [13]; [15] keeps contributions heading
    20: "",  # remove generic closing before literature review
    # --- LITERATURE REVIEW ---
    22: (
        "Early adversarial machine learning studies reported a puzzling inconsistency: networks that "
        "generalize well on held-out data can still assign high-confidence incorrect labels when "
        "inputs are modified within a small L\u221e budget. The failure is tied to how decision "
        "boundaries are arranged in high-dimensional pixel space, not to a single coding error in "
        "one layer.\n\n"
        "Goodfellow et al. proposed the Fast Gradient Sign Method (FGSM) as a practical one-step "
        "attack that follows the sign of the input gradient. Because FGSM requires only one forward "
        "and one backward pass per image, it is widely used in teaching labs and baseline robustness "
        "evaluations. Later work introduced stronger iterative and optimization-based attacks\u2014"
        "such as projected gradient descent and Carlini\u2013Wagner methods\u2014that often break "
        "defenses tuned only for single-step perturbations. Even so, FGSM remains a standard first "
        "test of whether a model reacts sensibly to gradient-directed input shifts."
    ),
    24: (
        "On the defense side, Madry et al. helped establish adversarial training as an empirical "
        "countermeasure: perturbed samples are generated during learning and optimized with their "
        "original labels. Reported results commonly show higher accuracy under attack alongside a "
        "modest reduction in clean-data performance. That accuracy\u2013robustness trade-off "
        "appears repeatedly across datasets and architectures.\n\n"
        "Parallel lines of AI security research examine poisoning of training data, model extraction "
        "through query access, membership inference on private records, and misuse of generative "
        "models. Together, these topics show that risk can appear during data collection, training, "
        "deployment, and inference. Against that background, hands-on reproduction remains useful. "
        "This study contributes an end-to-end implementation on a compact CIFAR-10 CNN, applying FGSM "
        "as the attack and adversarial training as the defense, with outcomes reported through "
        "measured accuracies, robustness curves, confusion matrices, and visual adversarial examples."
    ),
    25: "",
    # --- DISCUSSION ---
    197: (
        "The numbers from this implementation point to a clear split between clean-data competence "
        "and behavior under FGSM. The baseline CNN reached 72.54% on unperturbed CIFAR-10 test "
        "images, which is reasonable for a shallow architecture trained without augmentation. That "
        "same checkpoint was not reliable once FGSM was applied at test time. At \u03b5 = 0.03, "
        "accuracy fell to 7.34%\u2014below the 10% expected from random guessing across ten "
        "classes\u2014while the perturbed frames still looked nearly identical to the originals in "
        "side-by-side plots."
    ),
    198: (
        "From an implementation standpoint, the attack loop was straightforward: compute the input "
        "gradient, apply the signed update, clamp pixel values, and re-run inference. The ease of "
        "that procedure is itself part of the security concern. Any deployment where outside parties "
        "can submit images inherits exposure to gradient-directed tampering unless training "
        "explicitly accounts for it.\n\n"
        "Adversarial training changed the outcome without changing the network layout. The robust "
        "model dropped to 65.15% clean accuracy but recovered 28.38% attacked accuracy at "
        "\u03b5 = 0.03, compared with 7.34% for the baseline. Multi-epsilon curves showed the same "
        "pattern at 0.01, 0.05, and 0.10: lower clean scores, higher attacked scores. Confusion "
        "matrix output further showed that errors clustered among visually similar classes\u2014cat "
        "versus dog, automobile versus truck\u2014at 32\u00d732 resolution, which matches what the "
        "qualitative image pairs suggested during manual review."
    ),
    199: (
        "These observations do not imply full protection. FGSM was the only attack tested, the CNN "
        "is compact, and CIFAR-10 is a controlled benchmark. Even so, the measured gap between "
        "72.54% clean and 7.34% attacked accuracy on the baseline, and the partial recovery after "
        "retraining, support treating robustness as a property to engineer during model "
        "development\u2014especially for cyber security, autonomous perception, healthcare screening, "
        "and financial fraud pipelines where misclassification may carry operational cost."
    ),
    # --- CONCLUSION ---
    203: (
        "This project trained a CIFAR-10 CNN, stressed it with FGSM at four epsilon settings, and "
        "retrained a second model with adversarial examples included in each batch. The baseline "
        "reached 72.54% clean test accuracy but fell to 7.34% when attacked at \u03b5 = 0.03. "
        "After adversarial training, clean accuracy was 65.15%, while attacked accuracy at the same "
        "epsilon rose to 28.38%."
    ),
    204: (
        "The robustness plots, confusion matrix, and stored clean-versus-perturbed image pairs all "
        "reflected the same pattern: high clean accuracy alone did not predict stable behavior under "
        "gradient-based perturbation, and mixing adversarial samples into training partially closed "
        "that gap. For systems that may receive manipulated inputs in production, defense methods "
        "such as adversarial training should be considered during initial model design rather than "
        "added only after deployment."
    ),
}


def set_paragraph_text(paragraph, text: str) -> None:
    """Replace paragraph text while preserving the first run's formatting where possible."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    doc = Document(SRC)
    updated = 0
    for idx, new_text in REPLACEMENTS.items():
        if idx < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[idx], new_text)
            updated += 1
    doc.save(DST)
    print(f"Updated {updated} paragraphs in {DST}")


if __name__ == "__main__":
    main()
