"""
v8 - Complete replacement of ALL remaining old/AI paragraphs.
Uses exact text matching instead of index (indices shifted after deletions).
"""
from docx import Document

SRC = r'c:\Users\gaura\OneDrive\Documents\Adversarial Training\Adversaril attack and training Research paper_HUMANIZED.docx'
doc = Document(SRC)

def fix(para, new_text):
    runs = para.runs
    if not runs:
        return False
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''
    return True

# Old text → New text (exact match on old)
REPLACEMENTS = {
    # Duplicate/stale intro paragraph
    "The experiments follow a straightforward pipeline: build a baseline CNN on CIFAR-10, attack it with FGSM across several epsilon values, retrain it with adversarial training, and compare the results. Each step is described below.":
        "",

    # Duplicate dataset paragraph  
    "All experiments use CIFAR-10 as the image source. CIFAR-10 is a standard benchmark with 60,000 32×32 pixels distributed across 10 classes. The official split gives 50,000 images for training and 10,000 for testing.":
        "",

    # Duplicate "What I actually did"
    "What I actually did: built a CNN, trained it on CIFAR-10, attacked it with FGSM, watched it fail, retrained it with adversarial examples, and measured whether that helped. Accuracy tables, epsilon curves, confusion matrices, visual samples. All of it is in here.":
        "",

    # "The main libraries used were:" - still floating
    "The main libraries used were:":
        "Stack used:",

    # "y is the true class label." - orphaned
    "y is the true class label.":
        "",

    # "These metrics give complementary..." - duplicate
    "These metrics give complementary views: numbers quantify the damage, confusion matrices show which classes suffer most, and visual examples make the threat tangible.":
        "",

    # Duplicate "With a solid baseline in hand..."
    "With a solid baseline in hand, I ran FGSM at four epsilon values. The results are blunt.":
        "",

    # Duplicate numbers paragraph
    "Numbers tell the story directly: ε=0.01 → 23.13%, ε=0.03 → 7.34%, ε=0.05 → 3.35%, ε=0.10 → 1.04%. The model is performing below random chance at moderate attack strength.":
        "",

    # "On clean test data, the baseline CNN reached 72.54% accuracy." - duplicate
    "On clean test data, the baseline CNN reached 72.54% accuracy.":
        "",

    # "FGSM Attack Configuration" - stale section heading  
    "FGSM Attack Configuration":
        "",

    # "This accuracy-robustness tension isn't new..." - duplicate
    "This accuracy-robustness tension isn't new — the literature has documented it consistently, and these results land right where you'd expect.":
        "",

    # "Adversarial Example Analysis" - stale heading in wrong place
    "Adversarial Example Analysis":
        "",

    # "Cats and dogs blur together; automobiles and trucks do too..." - duplicate of [89]
    "Cats and dogs blur together; automobiles and trucks do too. That's expected — these pairs are genuinely similar, and adversarial noise amplifies the existing ambiguity. It suggests that targeted augmentation on these class pairs might help.":
        "",

    # "No catastrophic class-level failures either..." - duplicate
    "No catastrophic class-level failures either — the errors distribute in ways that make visual sense, rather than the model being effectively useless on specific categories.":
        "",

    # "Adversarial training changes the outcome..." - duplicate of [91]
    "Adversarial training changes the outcome. You give up around 7 percentage points of clean accuracy, and in return you get a model that scores 28.38% under the attack that broke the baseline. In security engineering, that's not a hard call.":
        "",

    # "If adversarial inputs are on the threat model..." - duplicate of [92]
    "If adversarial inputs are on the threat model — and for any security application they should be — then robustness needs to be a training objective from the start. Trying to add it after deployment is both harder and less effective.":
        "",

    # "Adversarial Training Performance" - stale section heading
    "Adversarial Training Performance":
        "",

    # "Security and accuracy pull in different directions..." - duplicate of [75]
    "Security and accuracy pull in different directions, but not incompatibly. Adversarial training shows you can have both — not perfectly, but well enough. The tools exist. The question is whether teams actually use them.":
        "",

    # "Deep learning models are powerful and brittle..." - stale old conclusion 
    "Deep learning models are powerful and brittle in equal measure. This paper demonstrates both properties concretely: a reasonable CNN collapses under adversarial pressure that no human would notice, and adversarial training brings it back to something workable. The 28.38% vs. 7.34% comparison at epsilon = 0.03 is the number that matters most — it shows that adversarial training isn't theoretical. It works, it ships, and it's worth the clean-accuracy cost.":
        "",

    # "Open questions from this work:" - duplicate
    "Open questions from this work:":
        "",

    # "A few directions worth pursuing next:" - duplicate
    "A few directions worth pursuing next:":
        "",

    # Old figure caption duplicate
    "Figure 6. FGSM Accuracy Comparison Between Standard CNN and Adversarial Trained CNN.":
        "Figure 6: Baseline vs. robust CNN accuracy across epsilon values.",

    # "Across all epsilon values, the robust model holds up far better than the baseline." - AI + duplicate
    "Across all epsilon values, the robust model holds up far better than the baseline.":
        "",

    # "For example: * Baseline CNN at ε = 0.03 → 7.34%"
    "For example: * Baseline CNN at ε = 0.03 → 7.34%":
        "",

    # "* Robust CNN at ε = 0.03 → 28.38%"
    "* Robust CNN at ε = 0.03 → 28.38%":
        "",

    # "The gap is substantial and validates adversarial training as a practical, meaningful defense."
    "The gap is substantial and validates adversarial training as a practical, meaningful defense.":
        "",

    # "Comparison Between Standard and Robust Models" - stale heading
    "Comparison Between Standard and Robust Models":
        "",

    # "The tradeoff is real but, in most security-oriented applications, well worth making..."
    "The tradeoff is real but, in most security-oriented applications, well worth making. The literature consistently reports this pattern, and the results here confirm it holds for this architecture and dataset.":
        "",

    # "Confusion Matrix Analysis" - stale section heading  
    "Confusion Matrix Analysis":
        "",

    # "I also generated a confusion matrix to get a class-by-class view of where errors concentrate."
    "I also generated a confusion matrix to get a class-by-class view of where errors concentrate.":
        "",

    # "Predictably, visually similar categories — cats vs. dogs..." - AI opener + duplicate
    "Predictably, visually similar categories — cats vs. dogs, automobiles vs. trucks — account for the bulk of off-diagonal entries. This kind of analysis helps pinpoint where the architecture could benefit most from targeted improvements.":
        "",

    # "It also confirms that the quantitative accuracy numbers are not hiding..."
    "It also confirms that the quantitative accuracy numbers are not hiding a model that performs well on some classes but fails catastrophically on others.":
        "",

    # "Discussion" - stale heading (there's already a results section)
    # (keep the Results and Discussion heading)

    # "The results cut through any doubt..." - now a duplicate
    "The results cut through any doubt. A model scoring 72.54% on clean data drops to 7.34% the instant adversarial perturbations enter — perturbations a human reviewer would never spot.":
        "",

    # "But the adversarial training results show there is a concrete path forward..." - AI phrase + duplicate
    "But the adversarial training results show there is a concrete path forward. Yes, clean accuracy drops a bit — but the model no longer collapses under attack, which is the more important property in a security context.":
        "",

    # "For any application where reliability under adversarial conditions matters..." - duplicate
    "For any application where reliability under adversarial conditions matters — security systems, medical imaging, autonomous vehicles — treating robustness as a core design goal rather than an afterthought is essential.":
        "",

    # "Limitations of the Study" - rename section heading
    "Limitations of the Study":
        "Limitations",

    # Old Limitations paragraph - stale
    "This work has real limitations worth being upfront about. The CNN used here is intentionally simple, and CIFAR-10, while standard, is not representative of every real-world distribution. More importantly, adversarial training against FGSM offers no guarantees against stronger attacks like PGD or Carlini-Wagner — evaluating those is a necessary next step. Future work should stress-test these defenses on deeper architectures, larger datasets, and a wider range of attack methods.":
        "",

    # Old conclusion - duplicate + AI phrases
    "AI's rapid rise has brought enormous benefits — and genuine security risks that can no longer be treated as edge cases. Adversarial examples sit near the top of that risk list, precisely because they expose a gap between how models and humans perceive the world. In this paper, I set out to quantify that vulnerability and test a practical remedy. The baseline CNN's accuracy cratered under even mild FGSM perturbations — from 72.54% down to single digits at epsilon = 0.03. Adversarial training brought meaningful improvement. The retrained model held 28.38% accuracy at epsilon = 0.03 versus the baseline's 7.34%, while clean-data accuracy stayed above 65%. Robustness curves, confusion matrices, and visual examples all told a consistent story.":
        "",

    # Junk text at end of conclusion
    "Ultimately, this work adds to a growing body of evidence that security needs to be built into the AI development process — and that adversarial training is a practical, deployable tool for doing exactly that.security-critical environments including cybersecurity, autonomous systems, and healthcare.":
        "",

    # "8.Future Work" - old heading
    "8.Future Work":
        "Future Work",

    # "Although the proposed defense mechanism improved robustness against FGSM attacks, several" - AI + cut off
    "Although the proposed defense mechanism improved robustness against FGSM attacks, several":
        "",

    # "opportunities remain for future research." - dangling
    "opportunities remain for future research.":
        "",

    # "First, the defense should be tested against stronger attacks..."
    "First, the defense should be tested against stronger attacks: PGD, DeepFool, and Carlini-Wagner are all natural next targets.":
        "",

    # "Second, comparing adversarial training against certified defenses..."
    "Second, comparing adversarial training against certified defenses and robust optimization would clarify when each approach is most appropriate.":
        "",

    # "Third, applying the same pipeline to ResNet or VGG-style architectures..."
    "Third, applying the same pipeline to ResNet or VGG-style architectures would test whether these findings generalize. 3.Evaluation of robustness on larger and more complex datasets.":
        "",

    # "4.Exploration of adversarial threats..."
    "4.Exploration of adversarial threats in domains such as cybersecurity, autonomous systems, healthcare, and privacy-preserving AI.":
        "",

    # "Each of these directions pushes toward the shared goal..."
    "Each of these directions pushes toward the shared goal: AI systems that are genuinely trustworthy, not just accurate under ideal conditions.":
        "",

    # "The visual comparison makes the point vividly..." - duplicate
    "The visual comparison makes the point vividly. Original and adversarial images look identical to a human. The CNN's prediction flips completely. In one test case, the model went from predicting class 3 to predicting class 5 — a completely different category — based on noise that no person would notice.":
        "",

    # "What comes next:" - duplicate
    "What comes next:":
        "",

    # "Immediate next steps from this work:" - duplicate
    "Immediate next steps from this work:":
        "",
}

changed = 0
for para in doc.paragraphs:
    txt = ''.join(r.text for r in para.runs).strip()
    if txt in REPLACEMENTS:
        new = REPLACEMENTS[txt]
        if new != txt:
            fix(para, new)
            changed += 1
            print(f"  Fixed: {txt[:60]}...")

print(f"\nFixed {changed} paragraphs")
doc.save(SRC)

# Final verify
doc2 = Document(SRC)
paras2 = [p for p in doc2.paragraphs if ''.join(r.text for r in p.runs).strip()]
print(f"Non-empty paragraphs: {len(paras2)}")

full = ' '.join(''.join(r.text for r in p.runs) for p in doc2.paragraphs)
print(f"Total chars: {len(full)}")

# Check for AI flags
flags = [
    'it is important','Furthermore,','Moreover,','Consequently,',
    'has been demonstrated','significantly improved','In conclusion,',
    'This paper aims','plays a crucial role','It is evident',
    'the proposed method','are vulnerable to','remain vulnerable',
    'This research investigates','The results clearly demonstrate',
    'Taken as a whole','Overall, the results',
    'Predictably,', 'This approach introduces',
]
bad = [f for f in flags if f.lower() in full.lower()]
print(f"AI flags: {'CLEAN' if not bad else bad}")
