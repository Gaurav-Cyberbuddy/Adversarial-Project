"""Insert Publication Enhancement Roadmap appendix into pdfff_formatted.docx."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

SRC = r"c:\Users\gaura\OneDrive\Documents\Adversarial Training\pdfff_formatted.docx"
DST = SRC


def find_paragraph(doc: Document, text: str) -> Paragraph | None:
    target = text.strip()
    for p in doc.paragraphs:
        if p.text.strip() == target:
            return p
    return None


def insert_paragraph_before(doc: Document, anchor: Paragraph, text: str, kind: str) -> Paragraph:
    new_p = anchor._element.makeelement(qn("w:p"), {})
    anchor._element.addprevious(new_p)
    para = Paragraph(new_p, doc)
    run = para.add_run(text)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if kind == "main":
        run.bold = True
        run.font.size = Pt(14)
    elif kind == "sub":
        run.bold = True
        run.font.size = Pt(12)
    elif kind == "list":
        para.paragraph_format.left_indent = Pt(18)
    else:
        run.font.size = Pt(11)
    return para


def main() -> None:
    doc = Document(SRC)
    anchor = find_paragraph(doc, "9. References")
    if anchor is None:
        raise SystemExit("Could not find References section.")

    blocks = [
        ("main", "Appendix A: Publication Enhancement Roadmap"),
        (
            "body",
            "Reviewer assessment (novelty and contribution). The completed experiments are "
            "technically sound and reproducible. The weakest publication aspect is limited "
            "research novelty: the pipeline reproduces a well-known FGSM attack and FGSM-based "
            "adversarial training on a shallow CNN using CIFAR-10. That combination is valuable "
            "as a student implementation report but is below the bar for most IEEE conference "
            "papers because the threat model (FGSM-only) matches the defense recipe (FGSM "
            "training), the architecture is small, and no cross-attack or cross-defense "
            "benchmarking is reported.",
        ),
        (
            "body",
            "The strongest existing assets are the end-to-end codebase, logged epsilon sweeps, "
            "confusion matrix, visual adversarial pairs, and the clear accuracy–robustness "
            "trade-off already measured at 72.54%, 7.34%, 65.15%, and 28.38%. The extension "
            "should build on these assets rather than replace them.",
        ),
        ("sub", "A. Three Realistic Extensions (Ranked)"),
        (
            "body",
            "Extension 1 — FGSM vs PGD robustness benchmarking (recommended). Evaluate baseline "
            "and FGSM-trained models under both FGSM and Projected Gradient Descent (PGD) at "
            "matched epsilon budgets. Report whether FGSM-trained robustness transfers to a "
            "stronger iterative attack.",
        ),
        (
            "list",
            "Publication impact: High — addresses the main limitation named in Section 8 and "
            "tests threat-model mismatch (gradient obfuscation / overfitting to weak attacks).",
        ),
        ("list", "Implementation difficulty: Low–Medium — reuses existing evaluation loop."),
        ("list", "Time required: 2–4 days on a laptop GPU/CPU."),
        (
            "body",
            "Extension 2 — Multi-defense comparison on a fixed PGD evaluation. Keep the same "
            "CNN and dataset; compare (a) standard training, (b) FGSM adversarial training "
            "already implemented, and (c) PGD adversarial training with identical epochs and "
            "optimizer settings.",
        ),
        (
            "list",
            "Publication impact: Medium–High — turns the project into a defense-comparison study.",
        ),
        ("list", "Implementation difficulty: Medium — requires a second training script."),
        ("list", "Time required: 3–5 days including PGD training time."),
        (
            "body",
            "Extension 3 — Attack transferability analysis. Generate adversarial examples on "
            "the baseline model and test whether they fool the robust model (and vice versa) "
            "without querying the target model gradients.",
        ),
        (
            "list",
            "Publication impact: Medium — adds a security-relevant black-box angle aligned with "
            "Papernot et al.",
        ),
        ("list", "Implementation difficulty: Low — no new training required initially."),
        ("list", "Time required: 1–2 days."),
        ("sub", "B. Recommended Extension: FGSM vs PGD Benchmarking"),
        (
            "body",
            "Why this option wins: it is the smallest code delta with the largest novelty gain. "
            "Reviewers will ask whether FGSM-trained robustness survives PGD; answering that "
            "question with measured tables is a genuine contribution. It also reframes the paper "
            "from a reproduction exercise into an empirical robustness evaluation study.",
        ),
        ("sub", "B.1 New Files to Create"),
        ("list", "attacks/attack_utils.py — shared clamping, loss backward, and epsilon loop."),
        ("list", "attacks/pgd_attack.py — PGD-L∞ generator (k steps, step size α = ε/4)."),
        ("list", "evaluation/multi_attack_eval.py — runs FGSM and PGD on both checkpoints."),
        ("list", "results/plot_attack_comparison.py — grouped bar/line chart for all runs."),
        ("sub", "B.2 Code Changes in Existing Files"),
        (
            "list",
            "attacks/fgsm_attack.py — import generate_fgsm() from attack_utils; return "
            "perturbed tensor instead of duplicating backward logic.",
        ),
        (
            "list",
            "evaluation/robust_test.py — accept --attack {fgsm,pgd} and --epsilon flags; print "
            "CSV-friendly rows for table export.",
        ),
        (
            "list",
            "results/plot_results.py — add Figure 8: FGSM vs PGD accuracy for Standard and "
            "Robust CNN across ε ∈ {0.01, 0.03, 0.05, 0.10}.",
        ),
        ("sub", "B.3 Additional Tables and Figures (after experiments complete)"),
        (
            "list",
            "TABLE VI — Baseline CNN accuracy under FGSM and PGD (four epsilon rows × two attack "
            "columns).",
        ),
        (
            "list",
            "TABLE VII — Robust CNN accuracy under FGSM and PGD (same layout).",
        ),
        (
            "list",
            "TABLE VIII — Robustness gap: (Robust PGD accuracy − Baseline PGD accuracy) per ε.",
        ),
        ("list", "Figure 8. FGSM versus PGD accuracy comparison for both models."),
        ("sub", "B.4 How the Paper Should Be Updated (no fabricated numbers)"),
        (
            "list",
            "Title — add evaluation scope, e.g., “…FGSM and PGD Attacks…”.",
        ),
        (
            "list",
            "Abstract — keep all current FGSM results; add one sentence stating PGD evaluation "
            "was performed and whether FGSM-trained robustness held under PGD (fill after runs).",
        ),
        (
            "list",
            "Methodology — add subsection “Projected Gradient Descent (PGD) Attack” with k=10, "
            "α=ε/4, L∞ projection; state both attacks use identical ε budgets.",
        ),
        (
            "list",
            "Results — add subsection “Cross-Attack Robustness” with Tables VI–VIII and Figure 8.",
        ),
        (
            "list",
            "Discussion — interpret FGSM–PGD gap; if PGD accuracy << FGSM accuracy on the robust "
            "model, report gradient obfuscation risk honestly.",
        ),
        (
            "list",
            "Conclusion — claim broadens to “robustness depends on the evaluated threat model”; "
            "retain all existing FGSM conclusions unchanged.",
        ),
        ("sub", "C. Step-by-Step Implementation Roadmap"),
        ("list", "Step 1 — Refactor FGSM into attacks/attack_utils.py; verify FGSM numbers unchanged."),
        ("list", "Step 2 — Implement PGD in attacks/pgd_attack.py; unit-test on one image batch."),
        (
            "list",
            "Step 3 — Run evaluation/multi_attack_eval.py on cnn_model.pth and "
            "adversarial_model.pth; save results/attack_benchmark.csv.",
        ),
        ("list", "Step 4 — Build Tables VI–VIII and Figure 8 from the CSV only."),
        ("list", "Step 5 — Update Abstract, Methodology, Results, Discussion, and Conclusion text."),
        ("list", "Step 6 — Add Madry PGD citation if not already cited; keep all five references."),
        (
            "body",
            "Important constraint: do not insert PGD accuracy values into the paper until the "
            "scripts finish and the CSV is verified. All existing FGSM tables, figures, and "
            "percentages in Sections 1–8 remain unchanged.",
        ),
    ]

    for kind, text in reversed(blocks):
        insert_paragraph_before(doc, anchor, text, kind)

    doc.save(DST)
    print(f"Inserted {len(blocks)} roadmap paragraphs before References in {DST}")


if __name__ == "__main__":
    main()
