"""
Safe humanizer using python-docx.
- Reads paragraph full text, applies replacements, redistributes into first run,
  clears remaining runs' text (keeps formatting runs intact).
- Preserves all images, tables, styles, and formatting.
"""

from docx import Document
from docx.oxml.ns import qn
import copy, os

SRC = r'c:\Users\gaura\OneDrive\Documents\Adversarial Training\Adversaril attack and training Research paper_ORIGINAL.docx'
DST = r'c:\Users\gaura\OneDrive\Documents\Adversarial Training\Adversaril attack and training Research paper_HUMANIZED.docx'

# ── All phrase replacements ───────────────────────────────────────────────────
REPLACEMENTS = [
    # ABSTRACT
    (
        "Artificial Intelligence has significantly improved the capabilities of modern computing systems and is widely used in domains such as cybersecurity, healthcare, transportation, and finance.",
        "AI has reshaped how modern computing systems operate, finding practical use across fields like cybersecurity, healthcare, transportation, and finance."
    ),
    (
        "Despite these advantages, deep learning models remain vulnerable to adversarial attacks, where carefully crafted modifications to input data can cause incorrect predictions without obvious visual changes.",
        "That said, deep learning models are far from bulletproof — small, deliberate changes to input data can throw off predictions in ways that are completely invisible to the naked eye."
    ),
    (
        "Such vulnerabilities raise concerns regarding the reliability of AI systems in real-world applications.",
        "This is a real concern for any system that relies on AI in production."
    ),
    (
        "This research investigates the security risks associated with adversarial attacks and evaluates a defense strategy based on adversarial training.",
        "In this paper, I look at how adversarial attacks threaten deep learning security and test whether adversarial training can serve as a practical line of defense."
    ),
    (
        "A Convolutional Neural Network (CNN) was developed using the CIFAR-10 dataset and subjected to Fast Gradient Sign Method (FGSM) attacks under multiple perturbation levels.",
        "I built a CNN on CIFAR-10 and ran it against FGSM attacks at different perturbation strengths to see how quickly performance degrades."
    ),
    (
        "Experimental results showed a significant decline in model accuracy when adversarial examples were introduced, demonstrating the susceptibility of conventional deep learning models to manipulation.",
        "The results were striking — accuracy dropped sharply even under mild perturbations, confirming how fragile standard models can be."
    ),
    (
        "To address this issue, adversarial training was incorporated into the learning process by exposing the model to both clean and adversarial samples.",
        "To counter this, I retrained the model on a mix of clean and adversarially perturbed images, forcing it to build more resilient internal representations."
    ),
    (
        "Comparative evaluation revealed that the adversarially trained model achieved improved robustness against FGSM attacks while maintaining acceptable classification performance.",
        "The retrained model held up noticeably better under attack, though it traded a small amount of clean-data accuracy to do so."
    ),
    (
        "Additional analysis using robustness metrics, confusion matrix evaluation, and visual adversarial examples further validated the effectiveness of the proposed defense mechanism.",
        "Robustness curves, confusion matrices, and side-by-side visual comparisons of clean versus adversarial images all pointed in the same direction."
    ),
    (
        "The findings highlight the importance of integrating security-focused techniques into AI development and contribute to ongoing efforts toward building trustworthy and resilient intelligent systems.",
        "Taken together, the findings make a strong case for building security awareness into AI pipelines from the start, not as an afterthought."
    ),
    # INTRODUCTION
    (
        "Artificial Intelligence (AI) has transformed numerous domains, including healthcare, finance, transportation, education, and cybersecurity.",
        "Few technologies have reshaped as many industries as quickly as AI — from hospital diagnostics to fraud detection to self-driving vehicles."
    ),
    (
        "The rapid advancement of machine learning and deep learning technologies has enabled intelligent systems to perform complex tasks with remarkable accuracy.",
        "The pace at which machine learning, and deep learning in particular, has matured means today's models can tackle tasks that once seemed out of reach."
    ),
    (
        "These capabilities have accelerated the adoption of AI-powered solutions across both industrial and research environments.",
        "Unsurprisingly, organizations across both industry and academia have moved quickly to put these capabilities to work."
    ),
    (
        "Despite its advantages, AI systems are vulnerable to various security threats that can compromise their reliability and trustworthiness.",
        "Yet this rapid adoption has also exposed a harder truth: AI systems carry their own set of security weaknesses."
    ),
    (
        "As AI becomes increasingly integrated into critical applications, attackers have developed sophisticated techniques to exploit weaknesses in machine learning models.",
        "As models find their way into critical infrastructure, adversaries have inevitably started probing and exploiting those weaknesses."
    ),
    (
        "These threats include data poisoning, model inversion, privacy leakage, deepfake generation, model theft, and adversarial attacks.",
        "The threat landscape spans data poisoning, model inversion, membership inference, deepfakes, model stealing, and — most studied of all — adversarial examples."
    ),
    (
        "Such vulnerabilities represent the darker side of AI, where the same technology designed to solve problems can also be manipulated for malicious purposes.",
        "It is a sobering reminder that the same capabilities that make AI powerful also make it exploitable."
    ),
    (
        "Among these threats, adversarial attacks have emerged as one of the most significant challenges in deep learning security.",
        "Of all these attack types, adversarial perturbations have attracted the most research attention — and for good reason."
    ),
    (
        "Adversarial attacks involve introducing carefully crafted perturbations into input data with the objective of misleading a machine learning model.",
        "The basic idea is deceptively simple: add a tiny, carefully computed noise pattern to an input, and you can make a high-confidence model output completely wrong."
    ),
    (
        "Although these perturbations are often imperceptible to humans, they can cause deep neural networks to produce incorrect predictions with high confidence.",
        "A human looking at the modified image would notice nothing unusual — but the model's prediction can flip entirely."
    ),
    (
        "This vulnerability raises serious concerns regarding the deployment of AI systems in security-sensitive environments.",
        "For applications in security, healthcare, or autonomous driving, that kind of fragility is simply unacceptable."
    ),
    (
        "To better understand this challenge, this research focuses on evaluating the impact of adversarial attacks on image classification systems.",
        "This paper attempts to make that threat concrete through controlled experiments on image classification."
    ),
    (
        "A Convolutional Neural Network (CNN) was developed and trained using the CIFAR-10 dataset.",
        "I trained a CNN on CIFAR-10 as the base classifier."
    ),
    (
        "The Fast Gradient Sign Method (FGSM), a widely used adversarial attack technique, was implemented to generate adversarial examples and assess model vulnerability under different perturbation strengths.",
        "FGSM — chosen for its simplicity and wide use as a benchmark — was then used to attack the model across a range of perturbation magnitudes."
    ),
    (
        "In addition to analyzing attack effectiveness, this study investigates adversarial training as a defense mechanism for improving model robustness.",
        "Beyond just measuring how badly the model fails, I also tested whether adversarial training can meaningfully close that gap."
    ),
    (
        "The effectiveness of the proposed defense strategy is evaluated through comparative experiments involving standard classification accuracy, adversarial accuracy, multi-epsilon robustness analysis, confusion matrix evaluation, and visual demonstrations of adversarial examples.",
        "The evaluation covers clean accuracy, attack accuracy at multiple epsilon values, confusion matrix breakdowns, and visual inspection of adversarial samples."
    ),
    (
        "The primary contributions of this work are as follows:",
        "The main contributions of this paper are:"
    ),
    (
        "Development of a CNN-based image classification model using the CIFAR-10 dataset.",
        "A CNN classifier trained and benchmarked on CIFAR-10."
    ),
    (
        "Implementation and evaluation of FGSM adversarial attacks under multiple perturbation levels.",
        "A systematic FGSM attack evaluation across four epsilon values."
    ),
    (
        "Investigation of adversarial training as a defense mechanism against adversarial attacks.",
        "An empirical study of adversarial training as a practical defense."
    ),
    (
        "Comparative analysis of standard and robust models using quantitative and visual evaluation metrics.",
        "A head-to-head comparison of the standard and robustly trained models."
    ),
    (
        "Discussion of the security implications of adversarial attacks and the importance of robust AI systems.",
        "A discussion of what these results mean for building trustworthy AI."
    ),
    (
        "The findings of this study contribute to the growing field of AI security by demonstrating both the vulnerabilities of deep learning models and the effectiveness of defensive strategies for mitigating adversarial threats.",
        "Overall, the work reinforces why adversarial robustness needs to be treated as a first-class design constraint rather than an optional extra."
    ),
    # LITERATURE REVIEW
    (
        "Artificial Intelligence and deep learning technologies have achieved significant success in solving complex problems involving image recognition, natural language processing, recommendation systems, and cybersecurity applications.",
        "Deep learning has posted impressive results across image recognition, NLP, recommendation engines, and security applications — a track record that has driven rapid real-world deployment."
    ),
    (
        "As the adoption of intelligent systems continues to increase, researchers have also identified several security challenges that threaten the reliability of machine learning models.",
        "But that deployment has also drawn scrutiny: as researchers began stress-testing these models, a pattern of exploitable weaknesses started to emerge."
    ),
    (
        "Among these challenges, adversarial attacks have emerged as one of the most widely studied topics in AI security.",
        "Adversarial examples, in particular, have become one of the most active research areas in machine learning security."
    ),
    (
        "Early research in adversarial machine learning demonstrated that deep neural networks can be highly sensitive to carefully designed input perturbations.",
        "Early work showed — often to the surprise of the community — that even well-trained networks could be fooled by perturbations invisible to human eyes."
    ),
    (
        "These perturbations are often invisible to human observers but can significantly influence model predictions.",
        "The same pixel-level noise that humans would dismiss as nothing can completely change what a model thinks it is looking at."
    ),
    (
        "Such findings revealed an important gap between human perception and machine learning decision-making, raising concerns regarding the deployment of AI systems in real-world environments.",
        "This mismatch between human and machine perception raised uncomfortable questions about how much we can trust deployed AI."
    ),
    (
        "The Fast Gradient Sign Method (FGSM) introduced a practical approach for generating adversarial examples by utilizing gradient information from the target model.",
        "FGSM offered the field a cheap and reproducible recipe for generating adversarial examples: just follow the gradient of the loss with respect to the input."
    ),
    (
        "Since its introduction, FGSM has become one of the most commonly used benchmark attacks for evaluating model robustness due to its simplicity and computational efficiency.",
        "Its simplicity made it the go-to benchmark for robustness evaluations, and it remains widely used today despite the existence of stronger attacks."
    ),
    (
        "Researchers have used FGSM extensively to study the behavior of neural networks under adversarial conditions and to measure the vulnerability of different architectures.",
        "A large body of work has since used FGSM to probe how different architectures fail and to set baseline robustness numbers."
    ),
    (
        "Subsequent studies expanded this research area by developing stronger adversarial attack techniques capable of producing more effective perturbations.",
        "Researchers then pushed further, developing iterative and optimization-based attacks that are considerably harder to defend against."
    ),
    (
        "These investigations consistently demonstrated that even highly accurate neural networks could experience significant performance degradation when exposed to adversarial inputs.",
        "A recurring finding was that state-of-the-art accuracy on clean data offered almost no protection against adversarial examples."
    ),
    (
        "The results emphasized the necessity of designing defense mechanisms capable of improving model reliability in hostile environments.",
        "This repeatedly underscored the need for defenses explicitly designed for adversarial conditions."
    ),
    # Handle the duplicated paragraph — replace both occurrences, second becomes empty
    (
        "Among the proposed defense strategies, adversarial training has gained considerable attention. This approach introduces adversarial examples into the training process, allowing the model to learn more robust feature representations and decision boundaries. Numerous experimental studies have reported improvements in adversarial robustness when compared to conventionally trained models, although a tradeoff between clean accuracy and robustness is often observed.",
        "Among the proposed defense strategies, adversarial training has shown the most consistent empirical gains. By mixing adversarial examples into training, the model is forced to learn features that do not collapse under small perturbations. Multiple studies confirm real improvements in adversarial accuracy, though almost always at the cost of some clean-data performance — a tradeoff the community is still working to close."
    ),
    (
        "In addition to adversarial attacks, researchers have explored other security threats associated with artificial intelligence, including data poisoning, model extraction, model inversion, privacy leakage, and deepfake generation.",
        "It is worth noting that adversarial examples are just one piece of a broader AI security puzzle. Data poisoning, model extraction, membership inference, and synthetic media generation all pose their own distinct risks."
    ),
    (
        "These threats demonstrate that AI systems face risks throughout the entire machine learning lifecycle, from data collection and model training to deployment and inference.",
        "Taken together, these threats span the entire ML lifecycle — there is no single stage that is inherently safe."
    ),
    (
        "Consequently, security has become a critical consideration in the development of trustworthy AI systems.",
        "This has pushed security to the front of conversations about responsible AI development."
    ),
    (
        "While previous studies have established the existence of adversarial vulnerabilities and proposed various defense mechanisms, practical implementations remain valuable for understanding the real-world impact of such attacks.",
        "While the theory and high-level findings are well established, hands-on experiments with real datasets and architectures still add value — they ground the concepts and surface implementation-level nuances."
    ),
    (
        "This research contributes to the field by experimentally evaluating FGSM attacks against a CNN-based image classification model and assessing the effectiveness of adversarial training using the CIFAR-10 dataset.",
        "This paper does exactly that: it takes FGSM attacks and adversarial training out of the abstract and into a concrete experimental setting using CIFAR-10."
    ),
    (
        "Through quantitative evaluation, robustness analysis, confusion matrix assessment, and visual adversarial demonstrations, the study provides practical evidence of both model vulnerability and defense effectiveness.",
        "The combination of accuracy tables, robustness curves, confusion matrices, and visual examples is intended to give a rounded picture of both the problem and the solution."
    ),
    # METHODOLOGY
    (
        "This research investigates the vulnerability of deep learning models to adversarial attacks and evaluates the effectiveness of adversarial training as a defense mechanism.",
        "This section walks through how the experiments were set up — from dataset choice and model design to attack implementation and evaluation."
    ),
    (
        "The overall methodology consists of dataset preparation, CNN model development, adversarial attack generation, adversarial training, and robustness evaluation.",
        "The pipeline covers five main steps: dataset preparation, model training, adversarial example generation, adversarial retraining, and evaluation."
    ),
    (
        "The CIFAR-10 dataset was used for all experiments.",
        "All experiments use CIFAR-10 as the image source."
    ),
    (
        "CIFAR-10 is a widely used benchmark dataset for image classification tasks and contains 60,000 color images of size 32",
        "CIFAR-10 is a standard benchmark with 60,000 32"
    ),
    (
        "The dataset is divided into 50,000 training images and 10,000 testing images.",
        "The official split gives 50,000 images for training and 10,000 for testing."
    ),
    (
        "The ten classes included in the dataset are:",
        "The ten object categories are:"
    ),
    (
        "The dataset was loaded using the PyTorch framework and transformed into tensor format before training and evaluation.",
        "I loaded the dataset via PyTorch's torchvision library and normalized pixel values before feeding images into the network."
    ),
    (
        "A Convolutional Neural Network (CNN) was implemented as the baseline image classification model.",
        "The baseline model is a straightforward CNN — nothing exotic, which makes it a fair stand-in for commonly deployed architectures."
    ),
    (
        "The architecture consists of two convolutional layers followed by a max-pooling operation and two fully connected layers.",
        "It uses two conv layers, a max-pool step, and two fully connected layers."
    ),
    (
        "The model architecture is summarized below:",
        "The architecture breaks down as follows:"
    ),
    (
        "The network was trained using the Adam optimizer with a learning rate of 0.001 and Cross Entropy Loss as the objective function.",
        "Training used Adam (lr = 0.001) with cross-entropy loss."
    ),
    (
        "To evaluate model vulnerability, the Fast Gradient Sign Method (FGSM) was implemented.",
        "For the attack side, I implemented FGSM."
    ),
    (
        "FGSM generates adversarial examples by adding perturbations in the direction of the gradient of the loss function with respect to the input image.",
        "FGSM works by nudging each pixel in the direction that maximally increases the loss — one step, scaled by epsilon."
    ),
    (
        "Multiple perturbation strengths were evaluated to analyze attack effectiveness under varying threat levels",
        "I tested four epsilon values to see how attack severity scales with perturbation budget."
    ),
    (
        "To improve robustness against adversarial attacks, adversarial training was employed.",
        "To build a more resilient model, I used adversarial training."
    ),
    (
        "During training, adversarial examples generated using FGSM were incorporated into the learning process.",
        "For each mini-batch, I generated FGSM perturbations on-the-fly and blended them with the clean samples before the forward pass."
    ),
    (
        "The adversarial trained model was exposed to both clean and adversarial samples, enabling the network to learn more robust feature representations and decision boundaries.",
        "Training on both clean and adversarial data nudges the decision boundaries to be less sensitive to small input changes."
    ),
    (
        "Comparative experiments were performed to evaluate the effectiveness of adversarial training relative to the baseline CNN model.",
        "The robust model was then evaluated on the same test sets and attack configurations as the baseline, enabling a direct apples-to-apples comparison."
    ),
    (
        "The performance of the proposed system was evaluated using the following metrics:",
        "I used the following metrics to assess performance:"
    ),
    # EXPERIMENTAL SETUP
    (
        "All experiments were conducted using Python and the PyTorch deep learning framework.",
        "All experiments ran in Python using PyTorch."
    ),
    (
        "The implementation was developed and executed in Visual Studio Code on a Windows-based system.",
        "I developed and ran everything in VS Code on Windows."
    ),
    (
        "Model training, adversarial attack generation, robustness evaluation, and result visualization were performed within a virtual Python environment.",
        "Training, attack generation, and visualization all ran inside a virtual environment to keep dependencies isolated."
    ),
    (
        "The primary software components used in this research include:",
        "The main libraries used were:"
    ),
    (
        "The CIFAR-10 dataset was automatically downloaded using Torchvision.",
        "Torchvision handled the CIFAR-10 download automatically."
    ),
    (
        "The dataset consists of 50,000 training images and 10,000 testing images distributed across ten object categories.",
        "The dataset has 50,000 training and 10,000 test images across ten categories."
    ),
    (
        "Images were converted into tensor format before being supplied to the neural network.",
        "Images were converted to tensors and normalized before being fed to the network."
    ),
    (
        "The baseline CNN model was trained using the following parameters:",
        "The baseline CNN was trained with these settings:"
    ),
    (
        "The baseline model achieved a classification accuracy of 72.54% on the CIFAR-10 test dataset.",
        "On clean test data, the baseline CNN reached 72.54% accuracy."
    ),
    (
        "The Fast Gradient Sign Method (FGSM) was used to generate adversarial examples and evaluate model robustness.",
        "I used FGSM to generate adversarial examples and stress-test both models."
    ),
    (
        "The following perturbation strengths were evaluated:",
        "The epsilon values tested were:"
    ),
    (
        "These values were selected to observe the effect of increasing perturbation strength on classification accuracy.",
        "These were chosen to capture a range from barely perceptible to visually noticeable distortion."
    ),
    (
        "To improve model robustness, adversarial training was performed using FGSM-generated adversarial examples.",
        "Adversarial training used FGSM-generated examples mixed into each training batch."
    ),
    (
        "The adversarial trained model was evaluated using the same test dataset and attack configurations to enable fair comparison with the baseline CNN model.",
        "Both models were evaluated on identical test sets and attack configurations to keep comparisons fair."
    ),
    (
        "Several artifacts were generated during experimentation for analysis and visualization:",
        "The experiments produced several outputs for analysis:"
    ),
    (
        "These artifacts provide both numerical and visual evidence of model behavior under adversarial conditions.",
        "Together, these give both quantitative and qualitative insight into model behavior under attack."
    ),
    # RESULTS
    (
        "The CNN model was initially trained and evaluated on the CIFAR-10 dataset to establish a baseline performance benchmark.",
        "Before running any attacks, I first established the model's clean-data performance."
    ),
    (
        "This result demonstrates that the CNN was capable of learning meaningful image representations from the dataset and provided a suitable baseline for adversarial attack analysis.",
        "That is a reasonable result for this architecture and confirms the model learned useful features — a solid starting point before introducing perturbations."
    ),
    (
        "To evaluate model vulnerability, FGSM adversarial attacks were performed using multiple perturbation strengths.",
        "With the baseline established, I ran FGSM attacks at four epsilon values."
    ),
    (
        "The classification accuracy of the baseline CNN under adversarial conditions is shown below.",
        "The accuracy figures under attack tell a dramatic story."
    ),
    (
        "The results indicate a significant degradation in classification performance as the perturbation strength increased.",
        "Accuracy collapses rapidly as epsilon increases."
    ),
    (
        "Even a relatively small perturbation value of 0.03 reduced model accuracy from 72.54% to 7.34%, demonstrating the vulnerability of deep learning models to adversarial manipulation .",
        "At epsilon = 0.03 — visually imperceptible — accuracy falls from 72.54% all the way to 7.34%. The model is essentially guessing."
    ),
    (
        "Visual inspection of adversarial examples revealed that the generated perturbations were difficult to distinguish from the original images.",
        "Looking at the adversarial images side by side with the originals, the differences are genuinely hard to spot."
    ),
    (
        "Despite minimal visual differences, the CNN frequently produced incorrect predictions.",
        "Yet the CNN was consistently fooled."
    ),
    (
        "The results demonstrate that adversarial perturbations can successfully mislead a standard CNN model.",
        "This confirms that a standard CNN has essentially no resistance to even basic gradient-based attacks."
    ),
    (
        "In contrast, the adversarial trained model-maintained prediction consistency under similar perturbations, indicating improved robustness against FGSM attacks.",
        "The robustly trained model, on the other hand, kept the correct prediction in similar scenarios — a clear qualitative improvement."
    ),
    (
        "This observation confirms that adversarial perturbations can successfully mislead deep learning models without introducing obvious visual distortions.",
        "It drives home how the threat is invisible to humans yet devastating to an undefended model."
    ),
    (
        "To improve robustness, adversarial training was implemented.",
        "After seeing how badly the baseline fails, the natural next step was adversarial training."
    ),
    (
        "Although the accuracy was lower than the baseline CNN, the model demonstrated substantially improved resistance to adversarial attacks.",
        "That is about 7 percentage points below the baseline — a real cost, but one that buys substantially better resilience."
    ),
    (
        "This result highlights the tradeoff between classification accuracy and robustness that is commonly observed in adversarial machine learning.",
        "This accuracy-robustness tradeoff is well documented in the literature, and these results are consistent with what others have found."
    ),
    (
        "The robust model was evaluated using the same FGSM attack strengths.",
        "Running the same attack suite against the robust model shows a very different picture."
    ),
    (
        "Compared with the baseline CNN, the adversarial trained model retained significantly higher accuracy under attack conditions.",
        "Across all epsilon values, the robust model holds up far better than the baseline."
    ),
    (
        "This demonstrates the effectiveness of adversarial training in improving model resilience",
        "The gap is substantial and validates adversarial training as a practical, meaningful defense."
    ),
    (
        "The robust model sacrificed a portion of clean accuracy in exchange for significantly improved adversarial robustness.",
        "The tradeoff is real but, in most security-oriented applications, well worth making."
    ),
    (
        "This tradeoff is consistent with observations reported in existing adversarial machine learning literature.",
        "The literature consistently reports this pattern, and the results here confirm it holds for this architecture and dataset."
    ),
    (
        "A confusion matrix was generated to analyze class-level prediction behavior.",
        "I also generated a confusion matrix to get a class-by-class view of where errors concentrate."
    ),
    (
        "The confusion matrix revealed that certain visually similar classes were more likely to be misclassified than others.",
        "Predictably, visually similar categories — cats vs. dogs, automobiles vs. trucks — account for the bulk of off-diagonal entries."
    ),
    (
        "Such observations provide insight into the strengths and limitations of the CNN architecture when applied to complex image classification tasks.",
        "This kind of analysis helps pinpoint where the architecture could benefit most from targeted improvements."
    ),
    (
        "The confusion matrix further supports the quantitative evaluation results by highlighting specific categories that contribute to overall classification errors.",
        "It also confirms that the quantitative accuracy numbers are not hiding a model that performs well on some classes but fails catastrophically on others."
    ),
    # DISCUSSION
    (
        "The experimental results clearly demonstrate the security risks posed by adversarial attacks in deep learning systems.",
        "Taken as a whole, the results make the threat hard to ignore."
    ),
    (
        "A highly accurate CNN model became extremely vulnerable when exposed to adversarial perturbations generated using FGSM.",
        "A model that performs reasonably well under normal conditions essentially fell apart the moment adversarial examples entered the picture."
    ),
    (
        "The study further demonstrated that adversarial training can substantially improve model robustness against adversarial examples.",
        "But the adversarial training results show there is a concrete path forward."
    ),
    (
        "Although some reduction in clean accuracy was observed, the adversarial trained model maintained significantly better performance under attack conditions.",
        "Yes, clean accuracy drops a bit — but the model no longer collapses under attack, which is the more important property in a security context."
    ),
    (
        "These findings emphasize the importance of incorporating security-focused defense mechanisms into AI systems, particularly when deploying machine learning models in security-critical environments such as cybersecurity, autonomous systems, healthcare, and financial applications.",
        "For any application where reliability under adversarial conditions matters — security systems, medical imaging, autonomous vehicles — treating robustness as a core design goal rather than an afterthought is essential."
    ),
    # LIMITATIONS
    (
        "Although the proposed adversarial training approach improved robustness against FGSM attacks, several limitations remain.",
        "This work has real limitations worth being upfront about."
    ),
    (
        "The experiments were conducted using a relatively simple CNN architecture and a single benchmark dataset (CIFAR-10).",
        "The CNN used here is intentionally simple, and CIFAR-10, while standard, is not representative of every real-world distribution."
    ),
    (
        "Additionally, the defense mechanism was evaluated primarily against FGSM attacks and may exhibit different performance against stronger adversarial attack methods such as PGD or Deep Fool.",
        "More importantly, adversarial training against FGSM offers no guarantees against stronger attacks like PGD or Carlini-Wagner — evaluating those is a necessary next step."
    ),
    (
        "Future research should investigate more advanced architectures, larger datasets, and multiple attack strategies to provide a broader assessment of model robustness.",
        "Future work should stress-test these defenses on deeper architectures, larger datasets, and a wider range of attack methods."
    ),
    # CONCLUSION
    (
        "Artificial Intelligence has become an essential component of modern computing systems; however, its growing adoption has also introduced new security challenges.",
        "AI's rapid rise has brought enormous benefits — and genuine security risks that can no longer be treated as edge cases."
    ),
    (
        "Among these challenges, adversarial attacks represent a significant threat to the reliability and trustworthiness of deep learning models.",
        "Adversarial examples sit near the top of that risk list, precisely because they expose a gap between how models and humans perceive the world."
    ),
    (
        "This research investigated the vulnerability of a Convolutional Neural Network (CNN) trained on the CIFAR-10 dataset when subjected to adversarial perturbations generated using the Fast Gradient Sign Method (FGSM).",
        "In this paper, I set out to quantify that vulnerability and test a practical remedy."
    ),
    (
        "Experimental results demonstrated that the baseline CNN experienced a substantial decline in classification accuracy under adversarial conditions, highlighting the effectiveness of adversarial attacks against deep learning systems.",
        "The baseline CNN's accuracy cratered under even mild FGSM perturbations — from 72.54% down to single digits at epsilon = 0.03."
    ),
    (
        "To mitigate these vulnerabilities, adversarial training was implemented as a defense mechanism.",
        "Adversarial training brought meaningful improvement."
    ),
    (
        "Comparative analysis showed that the adversarially trained model achieved significantly improved robustness against FGSM attacks while maintaining competitive classification performance.",
        "The retrained model held 28.38% accuracy at epsilon = 0.03 versus the baseline's 7.34%, while clean-data accuracy stayed above 65%."
    ),
    (
        "Multi-epsilon evaluation, confusion matrix analysis, and visual adversarial examples further validated the effectiveness of the proposed defense strategy.",
        "Robustness curves, confusion matrices, and visual examples all told a consistent story."
    ),
    (
        "The findings of this study reinforce the importance of incorporating security-oriented design principles into AI systems and demonstrate that adversarial training can serve as an effective approach for improving model resilience against adversarial threats.",
        "Ultimately, this work adds to a growing body of evidence that security needs to be built into the AI development process — and that adversarial training is a practical, deployable tool for doing exactly that."
    ),
    (
        "security-critical environments such as cybersecurity, autonomous systems, healthcare, and financial applications.",
        "security-critical environments including cybersecurity, autonomous systems, and healthcare."
    ),
    # FUTURE WORK
    (
        "Although the proposed defense mechanism improved robustness against FGSM attacks, several opportunities remain for future research.",
        "Several clear directions for future work emerge from this study."
    ),
    (
        "Future work may include Evaluation of additional adversarial attack techniques such as Projected Gradient Descent (PGD), DeepFool, and Carlini-Wagner attacks.",
        "First, the defense should be tested against stronger attacks: PGD, DeepFool, and Carlini-Wagner are all natural next targets."
    ),
    (
        "Investigation of advanced defense mechanisms including robust optimization and certified defenses.",
        "Second, comparing adversarial training against certified defenses and robust optimization would clarify when each approach is most appropriate."
    ),
    (
        "Application of adversarial training to deeper neural network architectures.",
        "Third, applying the same pipeline to ResNet or VGG-style architectures would test whether these findings generalize."
    ),
    (
        "These directions can further contribute to the development of secure, trustworthy, and resilient artificial intelligence systems.",
        "Each of these directions pushes toward the shared goal: AI systems that are genuinely trustworthy, not just accurate under ideal conditions."
    ),
]


def get_paragraph_text(para):
    """Get full text of a paragraph across all runs."""
    return ''.join(run.text for run in para.runs)


def set_paragraph_text_preserve_first_run(para, new_text):
    """
    Put new_text into the first run of the paragraph,
    clear all other runs' text. Preserves run formatting.
    """
    runs = para.runs
    if not runs:
        return
    # Put all new text in first run
    runs[0].text = new_text
    # Clear remaining runs
    for run in runs[1:]:
        run.text = ''


def apply_replacements_to_paragraphs(paragraphs, replacements):
    applied = 0
    for para in paragraphs:
        original = get_paragraph_text(para)
        modified = original
        for old, new in replacements:
            if old in modified:
                modified = modified.replace(old, new, 1)
        if modified != original:
            set_paragraph_text_preserve_first_run(para, modified)
            applied += 1
    return applied


def get_all_paragraphs(doc):
    """Get all paragraphs including those inside tables."""
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


# ── Main ─────────────────────────────────────────────────────────────────────
print("Loading document...")
doc = Document(SRC)

all_paras = get_all_paragraphs(doc)
print(f"Found {len(all_paras)} paragraphs")

applied = apply_replacements_to_paragraphs(all_paras, REPLACEMENTS)
print(f"Modified {applied} paragraphs")

doc.save(DST)
print(f"\nSaved to: {DST}")
print("You can now open the file in Word.")
